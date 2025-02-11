import base64
import shutil
import tempfile
import uuid
import threading
from django.contrib.sessions.models import Session
import threading
from bixdata_app.models import *

import pyperclip
from aiohttp.web_fileresponse import FileResponse
from django.core.files.storage import FileSystemStorage
from django.shortcuts import render, redirect
from django.http import JsonResponse, request
from django.contrib.auth import login, authenticate, logout
from django.contrib.auth.models import User
from django.core.mail import send_mail, BadHeaderError
from django.http import HttpResponse
from django.template.loader import render_to_string
import requests
import json
import datetime
from django.contrib.auth.decorators import login_required
import time
from datetime import timedelta

import pdfkit

import tempfile
from docx2pdf import convert as docx2pdf_convert

from .bixdata_view import *
from .businesslogic.office_calendar import OfficeCalendar
from ..forms import LoginForm
from django.contrib.auth.decorators import user_passes_test
from django.contrib import messages
from django.db import connection, transaction
from django.http import JsonResponse
from django.contrib.auth.models import Group, Permission, User, Group
from django_user_agents.utils import get_user_agent
import os
from django.conf import settings
from django.views.decorators.clickjacking import xframe_options_exempt
import subprocess
from .beta import *
from .helper_view import *
from htmldocx import HtmlToDocx
import csv
from functools import wraps
from .businesslogic.models.table import *
from .businesslogic.models.database_helper import *

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docxcompose.composer import Composer
from .helpers.helperdb import *

import qrcode

from .businesslogic.models.record import *
from .businesslogic.models.table import *

from django.middleware.csrf import get_token

from django.http import FileResponse, Http404


bixdata_server = os.environ.get('BIXDATA_SERVER')
freshdesk_apikey = os.environ.get('FRESHDESK_APIKEY')
get_tickets_password = os.environ.get('GET_TICKETS_PASSWORD')


# Questa funzione ritorna la pagina test_autocomplete.html
def get_test_autocomplete(request):
    return render(request, 'test_autocomplete.html')


# Questa funzione ritorna i dati richiesti con i select autocomplete
def get_autocomplete_data(request):
    term = request.GET.get('term')
    tableid = request.GET.get('tableid')
    recordid = request.GET.get('recordid')
    mastertableid = request.GET.get('mastertableid')
    recordidstabile = request.GET.get('recordidstabile')
    post = {
        'tableid': tableid,
        'recordid': recordid,
        'mastertableid': mastertableid,
        'recordidstabile': recordidstabile,
        'term': term
    }
    response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/get_autocomplete_data", data=post)

    response = json.loads(response.text)

    data = ['Red', 'Blue', 'Yellow', 'Green', 'Purple', 'Orange', '']
    data = [{"id": str(i), "value": item} for i, item in enumerate(
        data) if term.lower() in item.lower()]
    return JsonResponse({'data': response})


def get_full_data(request):
    if request.method == 'POST':
        name = request.POST.get('name')
        with connection.cursor() as cursor1:
            cursor1.execute(
                "SELECT query_conditions FROM sys_view WHERE id = 10"
            )
            query = cursor1.fetchone()[0]
            # query = query.replace('namee', name)
            with connection.cursor() as cursor2:
                cursor2.execute(
                    # "select dealname, expectedmargin, effectivemargin from user_deal where dealuser = %s", [name]
                    query
                )
                rows = cursor2.fetchall()
                dealname = [row[0] for row in rows]
                expectedmargin = [row[1] for row in rows]
                effectivemargin = [row[2] for row in rows]

                dealname = str(dealname)
                dealname = dealname.split(',')

                expectedmargin = str(expectedmargin)
                expectedmargin = expectedmargin.split(',')

                effectivemargin = str(effectivemargin)
                effectivemargin = effectivemargin.split(',')

                data2 = {
                    'dealname': dealname,
                    'expectedmargin': expectedmargin,
                    'effectivemargin': effectivemargin,
                    'name': name
                }

    return JsonResponse(data2)


def get_test_query2(request):
    with connection.cursor() as cursor1:
        cursor1.execute(
            "SELECT query_conditions FROM sys_view WHERE id = 11"
        )
        query = cursor1.fetchone()[0]
        with connection.cursor() as cursor2:
            cursor2.execute(
                query
            )
            rows = cursor2.fetchall()
            amounts = [row[0] for row in rows]
            months = [row[1] for row in rows]

            for i in range(len(amounts)):
                if amounts[i] is None:
                    amounts[i] = 0

            with connection.cursor() as cursor3:
                cursor3.execute(
                    "SELECT query_conditions FROM sys_view WHERE id = 15"
                )
                query2 = cursor3.fetchone()[0]
                with connection.cursor() as cursor4:
                    cursor4.execute(
                        query2
                    )
                    rows2 = cursor4.fetchall()
                    totalnets = [row[0] for row in rows2]

            data = {
                'amounts': amounts,
                'months': months,
                'totalnets': totalnets
            }

    return render(request, 'other/test_query2.html', {'data': data})


def get_full_data2(request):
    if request.method == 'POST':
        date = request.POST.get('date')
        with connection.cursor() as cursor1:
            cursor1.execute(
                "SELECT query_conditions FROM sys_view WHERE id = 12"
            )
            query = cursor1.fetchone()[0]
            with connection.cursor() as cursor2:
                cursor2.execute(
                    # "SELECT dealname, effectivemargin FROM user_deal WHERE closedate LIKE %s", [f"{date}-%"]
                    query
                )

                rows = cursor2.fetchall()
                dealname = [row[0] for row in rows]
                effectivemargin = [row[1] for row in rows]

                dealname = str(dealname)
                dealname = dealname.split(',')

                effectivemargin = str(effectivemargin)
                effectivemargin = effectivemargin.split(',')

                data2 = {
                    'dealname': dealname,
                    'effectivemargin': effectivemargin,
                    'date': date
                }

    return JsonResponse(data2)


def get_chart3(request):
    if request.method == 'POST':
        with connection.cursor() as cursor1:
            cursor1.execute(
                "SELECT query_conditions FROM sys_view WHERE id = 13"
            )
            query = cursor1.fetchone()[0]
            with connection.cursor() as cursor2:
                cursor2.execute(
                    query
                )
                rows = cursor2.fetchall()

                deal_users = [row[0] for row in rows]
                amounts = [row[1] for row in rows]

                data = {
                    'deal_users': deal_users,
                    'amounts': amounts,
                }
                with connection.cursor() as cursor3:
                    cursor3.execute(
                        "SELECT query_conditions FROM sys_view WHERE id = 14"
                    )
                    query2 = cursor3.fetchone()[0]
                    with connection.cursor() as cursor4:
                        cursor4.execute(
                            query2
                        )
                        rows2 = cursor4.fetchall()

                        amounts2 = [row[1] for row in rows2]
                        deal_users2 = [row[0] for row in rows2]

                        data2 = {
                            'amounts2': amounts2,
                            'deal_users2': deal_users2
                        }

        return render(request, 'other/chart3.html', {'data': data, 'data2': data2})


def get_chart4(request):
    if request.method == 'POST':
        with connection.cursor() as cursor1:
            cursor1.execute(
                "SELECT query_conditions FROM sys_view WHERE id = 16"
            )
            query = cursor1.fetchone()[0]
            with connection.cursor() as cursor2:
                cursor2.execute(
                    query
                )
                rows = cursor2.fetchall()

                months = [row[0] for row in rows]
                totalnets = [row[1] for row in rows]

                data = {
                    'months': months,
                    'totalnets': totalnets,
                }

        return render(request, 'other/chart4.html', {'data': data})


# Questa funzione ritorna il loading
def get_render_loading(request):
    theme = get_user_setting(request, 'theme')
    return render(request, 'other/loading.html', {'theme': theme})


# request: {tableid}


# Questa funzione serve per la ricerca dei record
def get_search_fields_data(tableid):
    search_fields = []
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT f.* FROM sys_user_table_search_field AS s join sys_field as f on s.tableid=f.tableid and s.fieldid=f.fieldid  WHERE s.tableid='{tableid}'"
        )
        result = dictfetchall(cursor)
        if result:
            for field in result:
                field['fieldtype'] = field['fieldtypeid']  # Add fieldtype key to the field dictionary
                search_fields.append(field)
    return search_fields


# Questa funzione è un test per i grafici
@login_required(login_url='/login/')
def get_render_content_chart(request):
    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT dealuser, SUM(effectivemargin), sum(expectedmargin) FROM user_deal WHERE dealuser is not null GROUP BY dealuser")
        rows = cursor.fetchall()
        users = [row[0] for row in rows]
        effective_margins = [row[1] for row in rows]
        expected_margins = [row[2] for row in rows]

        data = {
            'users': users,
            'effective_margins': effective_margins,
            'expected_margins': expected_margins,
        }

    return render(request, 'records/records_chart.html', {'data': data})


@login_required(login_url='/login/')
def get_block_records_chart(request):
    context = dict()
    tableid = request.POST.get('tableid')
    post = {
        'table': 'deal',
        'searchTerm': '',
    }
    response = requests.post(
        f"{bixdata_server}bixdata/index.php/rest_controller/get_records_chart", data=post)
    response_dict = json.loads(response.text)
    label = response_dict['label']

    label = '# of test'
    labels = ['Red', 'Blue', 'Yellow', 'Green', 'Purple', 'Orange']
    labels = json.dumps(labels)
    data = [12, 19, 3, 5, 2, 3]
    data = json.dumps(data)

    chart_data = list(zip(labels, data))

    records_table = render_to_string(
        'block/records/records_chart.html', chart_data, request=request)
    return HttpResponse(records_table)


# Questa funzione builda i dashboard blocks e li mette insieme nella pagina dashboard che ritorna
@login_required(login_url='/login/')
def get_render_content_dashboard(request):
    dbh = DatabaseHelper()
    context = {}
    context['blocks'] = []  # Initialize the blocks list
    context['block_list'] = []  # Initialize the block_list list
    user_id = request.user.id
    dashboard_id = request.POST.get('dashboard_id')

    with connection.cursor() as cursor2:

        cursor2.execute(
            "SELECT sys_user_id FROM v_users WHERE id = %s", [user_id]
        )
        bixid = cursor2.fetchone()[0]

        cursor2.execute(
            "SELECT dashboardid FROM sys_user_dashboard WHERE userid = %s", [bixid]
        )

        righe = cursor2.fetchall()
        # dashboard_id = righe[0][0]
        context['dashboardid'] = dashboard_id

    if request.method == 'POST':
        selected = ''
        with connection.cursor() as cursor:

            context['userid'] = bixid

            size = request.POST.get('size')
            context['size'] = size

            datas = SysUserDashboardBlock.objects.filter(userid=bixid, size=size, dashboardid=dashboard_id).values()

            # all_blocks = SysDashboardBlock.objects.all()
            sql = "SELECT * FROM sys_dashboard_block ORDER BY name asc"
            all_blocks = dbh.sql_query(sql)

            for block in all_blocks:
                context['block_list'].append(block)

            for data in datas:
                cursor.execute(
                    "SELECT * FROM v_sys_dashboard_block WHERE id = %s", [data['dashboard_block_id']]
                )
                results = dictfetchall(cursor)
                results = results[0]
                block = dict()
                block['id'] = data['id']

                block['gsx'] = data['gsx']
                block['gsy'] = data['gsy']
                block['gsw'] = data['gsw']
                block['gsh'] = data['gsh']
                block['viewid'] = results['viewid']
                block['widgetid'] = results['widgetid']

                # if they are null set default values
                if block['gsw'] == None or block['gsw'] == '':
                    block['gsw'] = 3
                    block['gsh'] = 2

                width = results['width']
                if width == None or width == 0 or width == '':
                    width = 4

                height = results['height']
                if height == None or height == 0 or height == '':
                    height = '50%'

                block['width'] = width
                block['height'] = height
                if results['reportid'] is None or results['reportid'] == 0:

                    if results['widgetid'] is None:
                        tableid = results['tableid']
                        tableid = 'user_' + tableid

                        block['html'] = get_records_table(request, results['tableid'], None, None, '',
                                                          results['viewid'], 1,
                                                          '', '')
                    else:
                        block['html'] = render_to_string('widgets/' + str(results['widgetid']) + '.html')


                else:

                    selected = ''
                    if results['operation'] == 'somma':
                        fields = results['fieldid'].split(';')
                        for field in fields:
                            field = 'SUM(' + field + ')'
                            selected += field + ','
                        groupby = results['groupby']
                        if results['custom'] == 'group_by_day':
                            groupby = f"DATE_FORMAT({groupby}, '%Y-%m-%d')"
                        if results['custom'] == 'group_by_month':
                            groupby = f"DATE_FORMAT({groupby}, '%Y-%m')"

                    query_conditions = results['query_conditions']
                    userid = get_userid(request.user.id)
                    query_conditions = query_conditions.replace("$userid$", str(userid))
                    id = data['id']
                    tableid = results['tableid']
                    name = results['name']
                    layout = results['layout']
                    fromtable = 'user_' + tableid
                    db = DatabaseHelper()
                    groupby_field_record = db.sql_query_row(
                        f"select * from sys_field where tableid='{tableid}' and fieldid='{results['groupby']}'")
                    if groupby_field_record['fieldtypeid'] == 'Utente':
                        fromtable = fromtable + f" LEFT JOIN sys_user ON {fromtable}.{results['groupby']}=sys_user.id "
                        selected += f"sys_user.firstname as {groupby}"
                    else:
                        selected += groupby

                    sql = "SELECT " + selected + " FROM " + fromtable + \
                          " WHERE " + query_conditions + " GROUP BY " + groupby
                    block['sql'] = sql
                    block['html'] = get_chart(request, sql, id, name, layout, fields)
                    context['userid'] = userid
                context['blocks'].append(block)

    return user_agent(request, 'content/dashboard.html', 'content/dashboard_mobile.html', context)


# Questa funzione serve per salvare le posizioni e le dimensioni dei blocchi della dashboard
def save_block_order(request):
    values = request.POST.get('value_list')

    values = json.loads(values)
    print(values)
    for value in values:
        record_id = value.get('id')
        size = value.get('size')
        gsx = value.get('gsX')
        gsy = value.get('gsY')
        gsw = value.get('gsW')
        gsh = value.get('gsH')

        if record_id is not None:
            SysUserDashboardBlock.objects.filter(id=record_id, size=size).update(gsx=gsx, gsy=gsy, gsw=gsw, gsh=gsh,
                                                                                 size=size)

    return JsonResponse({'success': True})


# Questa funzione prende in input i dati per comporre i grafici e ritorna il template del tipo di grafico in base alla variabile layout_chart
@login_required(login_url='/login/')
def get_chart(request, sql, id, name, layout, fields):
    query = sql
    id_sql = id
    name_chart = name
    layout_chart = layout
    fields_chart = fields

    with connection.cursor() as cursor2:
        cursor2.execute(query)
        rows = cursor2.fetchall()
        formatted_rows = []
        for row in rows:
            formatted_row = [str(value) if not isinstance(value, (int, float)) else value for value in row]
            formatted_rows.append(formatted_row)

        rows = formatted_rows
        value = []
        for num in range(0, len(fields_chart)):
            value.append([row[num] for row in rows])

        labels = [row[-1] for row in rows]

        if None in labels:
            labels = ['Non assegnato' if v is None else v for v in labels]

        for i in range(len(value)):
            for j in range(len(value[i])):
                if value[i][j] is not None:
                    if value[i][j] == 'None':
                        value[i][j] = 0
                    value[i][j] = round(value[i][j], 2)

        context = {
            'value': value,
            'labels': labels,
            'id': id_sql,
            'name': name_chart,
            'fields': fields_chart,
        }

        if layout_chart == 'barchart':
            return render_to_string('block/chart/barchart.html', context, request=request)
        elif layout_chart == 'piechart':
            return render_to_string('block/chart/piechart.html', context, request=request)
        elif layout_chart == 'linechart':
            return render_to_string('block/chart/linechart.html', context, request=request)
        elif layout_chart == 'horizontalbarchart':
            return render_to_string('block/chart/horizontalbarchart.html', context, request=request)
        elif layout_chart == 'donutchart':
            return render_to_string('block/chart/donutchart.html', context, request=request)


@login_required(login_url='/login/')
def record_card_duplicate(request):
    return render(request)


# Questa funzione serve per copiare la path del record
@login_required(login_url='/login/')
def record_card_copy(request):
    link = request.POST.get('link')
    pyperclip.copy(link)

    return True


# Questa funzione serve per elimiare un record in safe mode
@login_required(login_url='/login/')
def get_record_card_delete(request):
    if request.method == 'POST':
        recordid = request.POST.get('recordid')
        tableid = request.POST.get('tableid')

        with connection.cursor() as cursor:
            query = f"UPDATE user_{tableid} SET deleted_ = 'Y' WHERE recordid_= '{recordid}'"
            cursor.execute(
                query
            )
        deleted_record = Record(tableid=tableid, recordid=recordid)
        if tableid == 'dealline': 
            recordid_deal = deleted_record.fields['recordiddeal_'];
            custom_update(tableid='deal', recordid=recordid_deal)
    return JsonResponse({'success': True})


@user_passes_test(lambda u: u.is_superuser)
def get_render_gestione_utenti(request):
    return render(request, 'other/gestione_utenti.html')


# Questa funzione serve per buildare il gantt e ritornarlo in una pagina html (per ora non in utilizzo)
@login_required(login_url='/login/')
def get_block_records_gantt(request):
    context = dict()
    tableid = request.POST.get('tableid')
    post = {
        'table': 'task',
        'searchTerm': '',
    }

    response = requests.post(
        f"{bixdata_server}bixdata/index.php/rest_controller/get_records_gantt", data=post)
    response_dict = json.loads(response.text)

    records = response_dict['records']

    records_gantt = []
    for record in records:
        new_record = dict()
        new_record['id'] = record[1]
        new_record['name'] = record[2]
        new_record['start'] = record[3]
        new_record['end'] = record[4]
        new_record['progress'] = 100

        records_gantt.append(new_record)
    records_json = json.dumps(records_gantt)

    context = {
        'records': records,
        'records_json': records_json,
    }
    records_table = render_to_string(
        'block/records/records_gantt.html', context, request=request)
    return HttpResponse(records_table)


# Questa funzione serve per buildare il kanban e ritornarlo in una pagina html (per ora non in utilizzo)
@login_required(login_url='/login/')
def get_block_records_kanbanBAK(request):
    context = dict()
    tableid = request.POST.get('tableid')
    post = {
        'table': 'deal',
        'searchTerm': '',
    }

    response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/get_records_kanban", data=post)
    # response = requests.post("http://10.0.0.133:8822/bixdata/index.php/rest_controller/get_records_kanban", data=post)
    response_dict = json.loads(response.text)
    groups = response_dict['groups']
    return_groups = []
    for key, group in groups.items():
        return_group = dict()
        return_group['description'] = group['description']
        # da implementare il total
        return_group['totals'] = list()
        return_group['totals'].append({'totalname': 'totale 1', 'totalvalue': 100})
        return_group['totals'].append({'totalname': 'totale 2', 'totalvalue': 200})
        return_group_records = []
        return_record = dict()
        return_record['recordid'] = '123456789'
        return_record['title'] = 'title'
        return_record['tag'] = 'tag'
        return_record['date'] = 'date'
        return_record['user'] = 'user'
        return_record['field1'] = 'field1'
        return_record['field2'] = 'field2'
        return_record['field3'] = 'field3'
        return_record['field4'] = 'field4'
        return_group_records.append(return_record)
        return_group['records'] = group['records']
        return_groups.append(return_group)

    context = {
        'groups': return_groups,
        'tableid': tableid,
    }
    records_table = render_to_string(
        'block/records/records_kanban.html', context, request=request)
    return HttpResponse(records_table)


# Questa funzione serve per buildare il kanban e ritornarlo in una pagina html (per ora non in utilizzo)
@login_required(login_url='/login/')
def get_block_records_kanban(request):
    db = DatabaseHelper()
    context = dict()
    tableid = request.POST.get('tableid')
    searchTerm = request.POST.get('searchTerm')
    viewid = request.POST.get('viewid')
    filters = request.POST.get('filters')
    currentpage = request.POST.get('currentpage')
    order_field = request.POST.get('order_field')
    order = request.POST.get('order')
    fieldid = request.POST.get('fieldid')
    # record risultati

    # valori di raggruppamento (da rendere poi dinamico)
    if fieldid == '':
        sql = f"SELECT * FROM sys_lookup_table_item WHERE lookuptableid='dealstage_deal' order by itemorder asc"

    else:

        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT lookuptableid FROM sys_field WHERE fieldid = '{fieldid}' and tableid = '{tableid}'"
            )
            lookuptableid = dictfetchall(cursor)
            lookuptableid = lookuptableid[0]['lookuptableid']

        sql = f"SELECT * FROM sys_lookup_table_item WHERE lookuptableid='{lookuptableid}' order by itemorder asc"

    groups = db.sql_query(sql)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT sys_field.fieldid FROM sys_field JOIN sys_user_field_order ON sys_field.id = sys_user_field_order.fieldid WHERE sys_user_field_order.typepreference = 'kanban_fields'"
        )
        fieldids = dictfetchall(cursor)
        fieldids = [field['fieldid'] for field in fieldids]

        fields_str = ', '.join(fieldids)

    # totali per ogni gruppo

    # organizzazione valori di ritorno
    return_groups = []
    for group in groups:
        return_group = dict()
        dealstage = group['itemcode']
        return_group['description'] = group['itemdesc']
        # da implementare il total
        return_group['totals'] = list()
        return_group['totals'].append({'totalname': 'totale 1', 'totalvalue': 100})
        return_group['totals'].append({'totalname': 'totale 2', 'totalvalue': 200})

        sql = f"SELECT value, settingid FROM sys_user_table_settings WHERE tableid='{tableid}' AND settingid LIKE 'kanban_%'"
        kanban_settings = db.sql_query(sql)

        # separate settingids from values
        settings = []
        for setting in kanban_settings:
            set = setting['settingid'].split('_')
            setting['settingid'] = set[1]
            settings.append(setting)

        # now create a list of strings with values + 'AS' + settingid
        settings_str = ''
        for setting in settings:
            settings_str += f"{setting['value']} AS {setting['settingid']}"
            settings_str += ', '
        settings_str = settings_str[:-2]

        sql = f"SELECT recordid_ as recordid, {settings_str} FROM user_{tableid}  WHERE dealstage='{dealstage}' AND deleted_='n' ORDER BY closedate desc"
        records = db.sql_query(sql)
        return_group['records'] = records
        sql_totals = f"SELECT ROUND(SUM(amount)) as totalamount, ROUND(SUM(expectedmargin)) as totalmargin FROM user_deal WHERE dealstage='{dealstage}' AND deleted_='n'"
        totals = db.sql_query_row(sql_totals)
        return_group['totalamount'] = "{:,.2f}".format(totals['totalamount']).replace(",", "'") if totals[
                                                                                                       'totalamount'] is not None else "0.00"
        return_group['totalmargin'] = "{:,.2f}".format(totals['totalmargin']).replace(",", "'") if totals[
                                                                                                       'totalmargin'] is not None else "0.00"
        return_groups.append(return_group)

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT fieldid FROM sys_field WHERE (fieldtypeid = 'Utente' or fieldtypeid = 'Parola') and (lookuptableid is not NULL and lookuptableid != '') and tableid = '{tableid}'"
        )
        group_fields = dictfetchall(cursor)

    for group_field in group_fields:
        if group_field['fieldid'] == fieldid:
            group_field['selected'] = True

    context = {
        'groups': return_groups,
        'tableid': tableid,
        'group_fields': group_fields
    }
    records_table = render_to_string('block/records/records_kanban.html', context, request=request)
    return HttpResponse(records_table)


@login_required(login_url='/login/')
def get_block_records_calendar(request):
    tableid = request.POST.get('tableid')
    viewid = request.POST.get('viewid')
    datetype = request.POST.get('datetype')

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM sys_field WHERE fieldtypeid = 'Data' and tableid = '{tableid}'"
        )
        result = dictfetchall(cursor)

    table_obj = Table('task')
    conditions_list = list()
    # questa parte delle condition la sistemo io Alessandro, ma ho bisogno che mi arrivi qua la viewid
    userid = get_userid(request.user.id)
    conditions_list.append(f"user = {userid}")
    conditions_list.append("( status not like 'Chiuso' OR status is null)")
    conditions_list.append("duedate is not null")
    # conditions_list.append("id=832 OR id=863")

    if datetype is None:
        fieldid_date = result[1]['fieldid']
        datetype = fieldid_date + ' as date'
    else:
        fieldid_date = str(datetype)
        datetype = fieldid_date + ' as date'

    select_fields = ['recordid_', 'description as description', datetype, 'start as start', 'end as end']
    events_bixdata = table_obj.get_records(conditions_list=conditions_list, fields=select_fields)

    # Sostituisci i caratteri di nuova riga con <br> nella descrizione
    for event in events_bixdata:
        event['description'] = event['description'].replace('\r', ' ')
        event['description'] = event['description'].replace('\n', ' ')

    for option in result:
        if option['fieldid'] == fieldid_date:
            option['selected'] = 'selected'
        else:
            option['selected'] = ''

    return render(request, 'block/records/records_calendar.html',
                  {'events': events_bixdata, 'select_fields': result, 'tableid': tableid})


# Questa funzione non viene usata penso, quella nuova è in bixdata_view
@login_required(login_url='/login/')
def get_block_record(request):
    table = request.POST.get('table')
    searchTerm = request.POST.get('searchTerm')
    viewid = request.POST.get('viewid')
    post = {
        'table': table,
        'searchTerm': searchTerm,
        'viewid': viewid,
    }
    response = requests.post(
        f"{bixdata_server}bixdata/index.php/rest_controller/get_records", data=post)
    response_dict = json.loads(response.text)
    columns = response_dict['columns']
    records = response_dict['records']
    context = {
        'records': records,
        'columns': columns,
        'tableid': table,

    }

    for records_index, record in enumerate(records):
        for record_index, value in enumerate(record):
            if record_index == 4:
                # record[index]=split(value)
                value = value.split('|:|')
                record[record_index] = value[0]
            else:
                record[record_index] = value
        records[records_index] = record

    records_table = render_to_string(
        'block/records/records_table.html', context, request=request)
    return records_table


# Questa funzione richiama la funzione per creare la record card e ritorna la card come http response
@login_required(login_url='/login/')
def request_block_record_card(request):
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    master_tableid = request.POST.get('master_tableid')
    master_recordid = request.POST.get('master_recordid')
    userid = request.user.id
    return HttpResponse(get_block_record_card(request, tableid, recordid, userid, master_tableid, master_recordid))


# Questa funzione serve per creare la record card e ritorna la card come stringa
def get_block_record_card(request, tableid, recordid, userid, master_tableid='', master_recordid=''):
    context = dict()

    if tableid == 'salespush':
        with connection.cursor() as cursor:
            cursor.execute(f"SELECT status FROM user_salespush WHERE recordid_={recordid}")
            status = cursor.fetchone()[0]
            if status == 'Chiuso':
                display = 'none'
            else:
                display = 'block'

            context['display'] = display

    if tableid == 'timesheet':
        with connection.cursor() as cursor:
            cursor.execute(f"SELECT validated FROM user_timesheet WHERE recordid_={recordid}")
            validated = cursor.fetchone()[0]
            if validated == 'Si':
                display = 'none'
            else:
                display = 'block'

            context['display'] = display

    context['block_record_badge'] = get_block_record_badge(tableid, recordid)
    context['block_record_linked'] = get_block_record_linked(tableid, recordid)
    context['block_record_fields'] = ""
    context['recordid'] = recordid
    context['tableid'] = tableid
    context['master_tableid'] = master_tableid
    context['master_recordid'] = master_recordid
    context['layout_setting'] = get_user_setting(request, 'record_open_layout')
    if tableid == 'ticket':
        with connection.cursor() as cursor:
            cursor.execute(f"SELECT freshdeskid FROM user_ticket WHERE recordid_={recordid}")
            freshdeskid = cursor.fetchone()[0]
            context['freshdeskid'] = freshdeskid
    if tableid == 'deal':
        deal_record = Record(tableid='deal', recordid=recordid)
        context['dealstatus'] = deal_record.fields['dealstatus']
    if tableid == 'salesorder':
        with connection.cursor() as cursor:
            cursor.execute(f"SELECT id_bexio FROM user_salesorder WHERE recordid_={recordid}")
            id_bexio = cursor.fetchone()[0]
            context['id_bexio'] = id_bexio

    context['userid'] = userid
    context['user_table_settings'] = get_user_table_settings(userid, tableid)
    # TODO: recuperare i dati dal table settings generico
    table_settings = get_user_table_settings(1, tableid)

    if not 'default_recordtab' in table_settings:
        table_settings['default_recordtab'] = 'Dati'

    context['recordtab'] = table_settings['default_recordtab']
    # returned = user_agent(request, 'block/record/record_card.html', 'block/record/record_card_mobile.html', context)
    return render_to_string('block/record/record_card.html', context)


# Questa funzione richiama la funzione per creare il badge e lo ritorna
@login_required(login_url='/login/')
def request_block_record_badge(request, http_response=False):
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    # return HttpResponse({'success': True})
    return HttpResponse(get_block_record_badge(tableid, recordid))


# Questa funzione serve per creare il badge
def get_block_record_badge(tableid, recordid):
    print(f'get_block_record_badge {tableid} {recordid}')
    context = dict()

    post = {
        'tableid': tableid,
        'recordid': recordid,
    }
    # response = requests.post("http://10.0.0.133:8822/bixdata/index.php/rest_controller/get_fissi", data=post)
    # response_dict = json.loads(response.text)
    sql = f"SELECT sys_field.* FROM sys_field join sys_user_order on sys_field.fieldid=sys_user_order.fieldid WHERE sys_user_order.userid=1 AND sys_user_order.tableid='{tableid}' AND typePreference='campiFissi' ORDER BY fieldorder asc"

    fields = db_query_sql(sql)
    values = db_get_row(f"user_{tableid}", "*", f"recordid_='{recordid}'")
    context_fields = dict()
    for field in fields:
        fieldid = field['fieldid']
        field['value'] = values[fieldid]
        context_fields[fieldid] = field

    context['fields'] = context_fields

    records_table = ""
    block_record_badge = ''

    if tableid == 'company':
        sql = f"SELECT DISTINCT type FROM user_servicecontract WHERE recordidcompany_='{recordid}' AND STATUS='In Progress'"
        context_fields['services'] = db_query_sql(sql)
        context['fields'] = context_fields
        block_record_badge = render_to_string('block/record/custom/record_badge_company.html', context)
    elif tableid == 'deal':
        deal_record = Record('deal', recordid)
        company_record = Record(tableid='company', recordid=deal_record.fields['recordidcompany_'])
        deal_record.fields['companyname'] = company_record.fields['companyname']
        context['fields'] = deal_record.fields
        block_record_badge = render_to_string('block/record/custom/record_badge_deal.html', context)
    else:
        block_record_badge = render_to_string('block/record/record_badge.html', context)
    return block_record_badge


@login_required(login_url='/login/s')
def get_block_record_fields(request, prefilled_fields=dict()):
    context = dict()
    http_response = request.POST.get('http_response')
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    ticketid = request.POST.get('ticketid')
    recordid_ticket = request.POST.get('recordid_ticket')
    master_tableid = request.POST.get('master_tableid')
    master_recordid = request.POST.get('master_recordid')
    contextfunction = request.POST.get('contextfunction')
    contextreference = request.POST.get('contextreference')
    # creator = request.POST.get('creator')


    row = SysUser.objects.filter(bixid=request.user.id).values('id')

    if row:
        userid = row[0]
        userid = userid['id']
        context['userid'] = userid

    record = Record(tableid=tableid, recordid=recordid, userid=userid)
    record.master_tableid = master_tableid
    record.master_recordid = master_recordid
    record.context = contextfunction




    fields = record.get_fields()


    for key, value in fields['Dati'].items():
        if value['fieldtypewebid'] == 'html':
            value['unique_id'] = str(uuid.uuid4().hex)

    #check se creator è lo stesso dello user che visualizza, in caso contrario, duedate non è modificabile
    if tableid == 'task' and contextfunction != 'insert':
        creator = fields['Dati']['creator']['valuecode'][0]['code']
        userid = str(get_userid(request.user.id))
        if creator != userid:
            fields['Dati']['duedate']['edit'] = 'disabled'
        else:
            fields['Dati']['duedate']['edit'] = ''

    context['record_fields_labels'] = fields
    context['contextfunction'] = contextfunction
    context['contextreference'] = contextreference
    context['tableid'] = tableid
    context['recordid'] = recordid
    context['master_tableid'] = master_tableid
    context['master_recordid'] = master_recordid
    


    bixuserid = get_userid(request.user.id)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT value FROM sys_user_settings WHERE setting = 'record_open_layout' AND userid = %s",
            [bixuserid]
        )
        user_setting = cursor.fetchone()
        user_setting = user_setting[0]

        context['user_setting'] = user_setting

    if tableid == 'timesheet':
        context['timesheet'] = uuid.uuid4()
        if contextfunction != 'insert':
            timesheet_record = Record(tableid='timesheet', recordid=recordid)

            if timesheet_record.fields['validated'] == 'Si':
                context['edit_block'] = True
            else:
                context['edit_block'] = False

            if userid in [53, 2, 47, 50, 3]:
                context['edit_block'] = False
        else:
            context['edit_block'] = False

        context['block_record_fields_timesheet'] = render_to_string('block/record/record_fields.html', context,
                                                                    request=request)
        context['block_record_fields'] = render_to_string('block/record/custom/record_fields_timesheet.html', context,
                                                          request=request)
        block_record_fields_container = render_to_string('block/record/record_fields_container.html', context,
                                                         request=request)
    else:


        context['block_record_fields'] = render_to_string('block/record/record_fields.html', context, request=request)
        block_record_fields_container = render_to_string('block/record/record_fields_container.html', context,
                                                         request=request)

    if (http_response):
        return HttpResponse(block_record_fields_container)
    else:
        return block_record_fields_container


def new_timesheet(request):
    timetr_recordid = request.POST.get('timetr_recordid')
    prefilled_fields = dict()
    prefilled_fields['description'] = 'testvalue'
    prefilled_fields['worktime'] = '01:30'

    block = get_block_record_fields(request, prefilled_fields=prefilled_fields)
    return HttpResponse(block)


@login_required(login_url='/login/')
def get_block_record_linked_OLD(request):
    context = dict()
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    post = {
        'tableid': tableid,
        'recordid': recordid,
    }
    response = requests.post(
        f"{bixdata_server}bixdata/index.php/rest_controller/get_record_labels", data=post)
    response_dict = json.loads(response.text)
    context['labels'] = response_dict
    context['tableid'] = tableid
    context['recordid'] = recordid
    record_linked_labels = render_to_string(
        'block/record/record_linked.html', context, request=request)
    return record_linked_labels


@login_required(login_url='/login/')
def request_block_record_linked(request):
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    return HttpResponse(get_block_record_linked(tableid, recordid))


def get_block_record_linked(tableid, recordid):
    context = dict()

    rows = db_query_sql(
        f"SELECT sys_table_link.*,typepreference,sys_user_order.ID FROM sys_table_link JOIN sys_user_order ON sys_table_link.tableid=sys_user_order.tableid AND sys_table_link.tablelinkid=sys_user_order.fieldid  WHERE sys_table_link.tableid = '{tableid}' AND typepreference='keylabel'")

    linked_tables = list()
    for row in rows:
        linked_table = dict()
        linked_tableid = row['tablelinkid']
        table_name = 'test'
        linked_table['table_name'] = db_get_value("sys_table", "description", f"id='{linked_tableid}'")
        table_count = db_get_count(f"user_{linked_tableid}", f"recordid{tableid}_='{recordid}'")
        linked_table['table_count'] = table_count
        linked_table['tableid'] = linked_tableid
        linked_tables.append(linked_table)

    context['linked_tables'] = linked_tables
    context['tableid'] = tableid
    context['recordid'] = recordid

    record_linked_labels = render_to_string('block/record/record_linked.html', context)
    return record_linked_labels


@login_required(login_url='/login/')
def get_linked(request):
    name = ''
    if request.method == 'POST':
        name = request.POST.get('name')
    return JsonResponse(name)


@login_required(login_url='/login/')
def save_record_fields(request):
    db_helper = DatabaseHelper('default')
    file = request.FILES.get('file')
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    contextfunction = request.POST.get('contextfunction')
    creator = userid = get_userid(request.user.id)
    post_dict = request.POST.dict()
    post_repr = repr(post_dict)
    post_repr = post_repr.replace("'", "")
    sql = f"INSERT INTO sys_logquery (userid,funzione,post) VALUES ({userid},'salva_record_fields','{post_repr}')"
    db_helper.sql_execute(sql)
    fields_dict = request.POST.dict()
    del fields_dict['tableid']
    del fields_dict['recordid']
    del fields_dict['contextfunction']

    if tableid == 'timetracking':
        if fields_dict['stato'] == 'Terminato':
            if fields_dict['end'] == '':
                fields_dict['end'] = datetime.datetime.now().strftime("%H:%M")
            time_format = '%H:%M'
            start = datetime.datetime.strptime(fields_dict['start'], time_format)
            end = datetime.datetime.strptime(fields_dict['end'], time_format)
            time_difference = end - start

            total_minutes = time_difference.total_seconds() / 60
            hours, minutes = divmod(total_minutes, 60)
            formatted_time = "{:02}:{:02}".format(int(hours), int(minutes))

            fields_dict['worktime_string'] = str(formatted_time)

            hours = time_difference.total_seconds() / 3600
            fields_dict['worktime'] = round(hours, 2)

        if fields_dict['start'] == '':
            fields_dict['start'] = datetime.datetime.now().strftime("%H:%M")

    selected_options = request.POST.getlist('service');

    # TODO: da testare e completare
    # if tableid == 'task':
    # oc = OfficeCalendar()
    # if fields_dict['planneddate'] != '':
    # if fields_dict['o365_idcalendar'] == '':
    # fields_dict['o365_idcalendar'] = oc.add_calendar_event(fields_dict)
    # else:
    # oc.update_calendar_event(fields_dict)

    fields = json.dumps(fields_dict)

    post_data = {
        'tableid': tableid,
        'recordid': recordid,
        'fields': fields
    }

    response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/set_record", data=post_data)
    response_dict = json.loads(response.text)




    if contextfunction == 'edit':
        if tableid == 'task':
            check_task_status(recordid)


    if tableid == 'ticketbixdata' and 'description' in fields_dict:
        message = 'Nuovo ticket aperto da {} \nDescrizione: {}\nTipo: {}'.format(request.user.username,
                                                                                 fields_dict['description'],
                                                                                 fields_dict.get('type', 'N/A'))
        send_email(emails=['marco.garganigo@swissbix.ch', 'alessandro.galli@swissbix.ch'],
                   subject='Supporto bixdata', html_message=message)


    elif tableid == 'task':


        if (contextfunction == 'edit' or contextfunction == 'insert') and fields_dict['completed'] != 'Si':

            creator = str(creator)

            if fields_dict['user'] != fields_dict['creator']:
                with connection.cursor() as cursor:
                    cursor.execute("SELECT email FROM v_users WHERE sys_user_id = %s", [fields_dict['user']])
                    row = cursor.fetchone()

                    email = row[0]

                    cursor.execute("SELECT first_name, last_name FROM v_users WHERE sys_user_id = %s",
                            [fields_dict['creator']])
                    row = cursor.fetchone()
                    first_name = row[0]
                    last_name = row[1]

                    companyname = ''
                    projectname = ''

                    if fields_dict['recordidcompany_'] != 'None':
                        with connection.cursor() as cursor:
                            cursor.execute("SELECT companyname FROM user_company WHERE recordid_ = %s",
                                           [fields_dict['recordidcompany_']])
                            row = cursor.fetchone()
                            companyname = row[0]

                    if fields_dict['recordidproject_'] != 'None':
                        with connection.cursor() as cursor:
                            cursor.execute("SELECT projectname FROM user_project WHERE recordid_ = %s",
                                           [fields_dict['recordidproject_']])
                            row = cursor.fetchone()
                            projectname = row[0]

                    fields_dict['username'] = first_name + ' ' + last_name
                    fields_dict['companyname'] = companyname
                    fields_dict['projectname'] = projectname

                    fields_dict['recordid'] = response_dict['recordid']

                    message = render_to_string('other/new_task.html', fields_dict)

                    score = ''

                    if companyname != '':
                        score = ' - '

                    # return render(request, 'other/new_task.html', fields_dict)

                    send_email(emails=[email],
                               subject='Nuovo task assegnato da ' + fields_dict['username'] + score + companyname,
                               html_message=message)

    elif tableid == 'salespush':

        creator = str(creator)

        if fields_dict['assignedto'] != creator:
            with connection.cursor() as cursor:
                cursor.execute("SELECT email FROM v_users WHERE sys_user_id = %s", [fields_dict['assignedto']])
                row = cursor.fetchone()

                email = row[0]

                cursor.execute("SELECT first_name, last_name FROM v_users WHERE sys_user_id = %s",
                               [creator])
                row = cursor.fetchone()
                first_name = row[0]
                last_name = row[1]

                if fields_dict['recordidcompany_'] != 'None':
                    with connection.cursor() as cursor:
                        cursor.execute("SELECT companyname FROM user_company WHERE recordid_ = %s",
                                       [fields_dict['recordidcompany_']])
                        row = cursor.fetchone()
                        companyname = row[0]

                    fields_dict['company'] = companyname
                fields_dict['description'] = fields_dict['description']
                fields_dict['username'] = first_name + ' ' + last_name
                fields_dict['recordid'] = response_dict['recordid']

                message = render_to_string('other/new_salespush.html', fields_dict)

                # return render(request, 'other/new_task.html', fields_dict)

                send_email(emails=[email], subject='Nuovo push commerciale', html_message=message)

    for field_name, uploaded_files in request.FILES.items():
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        attachments_path = os.path.join(base_dir, 'attachments')
        fs = FileSystemStorage(location=attachments_path)
        print("Location: "+fs.location)
        #fs_bix = FileSystemStorage(location='attachments_bixdata')
        #print("Location:"+fs_bix.location)
        basename, extension = os.path.splitext(uploaded_files.name)
        filename = tableid + '_' + response_dict['recordid']
        if tableid == 'attachment':
            if fields_dict['type'] == 'Documento firmato':
                filename = 'deal' + '_' + fields_dict['recordiddeal_']
            if fields_dict['type'] == 'Allegato generico':
                filename = 'deal-attachment' + '_' + fields_dict['recordiddeal_']
        current_datetime = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = filename + '_' + basename + '_' + current_datetime + extension

        #with connection.cursor() as cursor:
         #   cursor.execute(
          #      f"UPDATE user_attachment SET filename = '{filename}' WHERE recordid_ = '{response_dict['recordid']}'"
           # )

        fs.save(filename, uploaded_files)
        record = Record(tableid, response_dict['recordid'])
        record.fields[field_name]=filename
        record.save()
        #fs_bix.save(filename, uploaded_files)

    # return render(request, 'block/record/record_fields.html')
    custom_save_record(request, tableid, response_dict['recordid'])
    return HttpResponse(response_dict['recordid'])


@login_required(login_url='/login/')
def custom_save_record(request, tableid, recordid):
    # ---DEAL---
    if tableid == 'deal':
        record_deal = Record('deal', recordid)
        record_company = Record('company', record_deal.fields['recordidcompany_'])
        reference = str(record_deal.fields['id']) + ' - ' + record_company.fields['companyname'] + ' - ' + \
                    record_deal.fields['dealname']
        if record_deal.fields['advancepayment'] == None:
            record_deal.fields['advancepayment'] = 0
        record_deal.fields['reference'] = reference
        if record_deal.fields['dealstatus'] is None or (record_deal.fields['dealstatus'] != 'Vinta' and record_deal.fields['dealstatus'] != 'Persa'):
            record_deal.fields['dealstatus']='Aperta'
        record_deal.save()

    # ---TIMESHEET---
    if tableid == 'timesheet':
        # recupero informazioni necessarie
        timesheet_table = Table(tableid='timesheet')
        servicecontract_table = Table(tableid='servicecontract')
        timesheet_record = Record('timesheet', recordid)
        company_record = Record('company', timesheet_record.fields['recordidcompany_'])
        project_record = Record('project', timesheet_record.fields['recordidproject_'])
        ticket_record = Record('ticket', timesheet_record.fields['recordidticket_'])
        servicecontract_record = Record('servicecontract', timesheet_record.fields['recordidservicecontract_'])
        service = timesheet_record.fields['service']
        invoiceoption = timesheet_record.fields['invoiceoption']
        invoicestatus = timesheet_record.fields['invoicestatus']
        if isempty(invoicestatus):
            invoicestatus = ''
        worktime = timesheet_record.fields['worktime']
        traveltime = timesheet_record.fields['traveltime']

        # inizializzo campi
        productivity = ''
        worktime_decimal = 0
        travel_time_decimal = 0
        totaltime_decimal = 0
        workprice = 0
        travelprice = 0
        totalprice = 0
        timesheet_record.fields['worktime_decimal'] = ''
        timesheet_record.fields['traveltime_decimal'] = ''
        timesheet_record.fields['totaltime_decimal'] = ''
        timesheet_record.fields['workprice'] = ''
        timesheet_record.fields['travelprice'] = ''
        timesheet_record.fields['totalprice'] = ''
        timesheet_record.fields['recordidservicecontract_'] = '';
        timesheet_record.fields['print_type'] = 'Normale'
        timesheet_record.fields['print_hourprice'] = ''
        timesheet_record.fields['print_travel'] = ''
        # aggiorno dati a prescindere

        if not isempty(worktime):
            hours, minutes = map(int, worktime.split(':'))
            worktime_decimal = hours + minutes / 60
            if not isempty(traveltime):
                hours, minutes = map(int, traveltime.split(':'))
                travel_time_decimal = hours + minutes / 60
            totaltime_decimal = worktime_decimal + travel_time_decimal
            timesheet_record.fields['worktime_decimal'] = worktime_decimal
            timesheet_record.fields['traveltime_decimal'] = travel_time_decimal
            timesheet_record.fields['totaltime_decimal'] = totaltime_decimal

        # inizio valutazione invoice status

        # se ho già un service contract lo mantengo
        # if not isempty(servicecontract_record.recordid) and invoiceoption!='Out of contract':
        #   if servicecontract_record.fields['type']=='Monte Ore':
        #      timesheet_record.fields['recordidservicecontract_']=servicecontract_record.recordid
        #     invoicestatus="Service Contract: Monte Ore"
        #    productivity='Ricavo diretto'
        # else:
        #    if invoicestatus!='Invoiced':
        #       invoicestatus='To Process'
        # else:
        #   if invoicestatus!='Invoiced':
        #      invoicestatus='To Process'

        if invoicestatus != 'Invoiced':
            invoicestatus = 'To Process'

        # valutazione del tipo di servizio se produttivo o meno
        if invoicestatus == 'To Process':
            if service == 'Amministrazione' or service == 'Commerciale' or service == 'Formazione Apprendista' or service == 'Formazione e Test' or service == 'Interno' or service == 'Riunione':
                invoicestatus = 'Attività non fatturabile'
                productivity = 'Senza ricavo'

        # valutazione delle option
        if invoicestatus == 'To Process':
            if invoiceoption == 'Under Warranty' or invoiceoption == 'Commercial support' or invoiceoption == 'Swisscom incident' or invoiceoption == 'Swisscom ServiceNow' or invoiceoption == 'To check':
                invoicestatus = invoiceoption
                productivity = 'Senza ricavo'
                timesheet_record.fields['print_type'] = 'Garanzia'
                timesheet_record.fields['print_hourprice'] = 'Garanzia'
                timesheet_record.fields['print_travel'] = 'Garanzia'

        # valutazione eventuale project
        if invoicestatus == 'To Process' and (
                (not isempty(project_record.recordid)) and invoiceoption != 'Out of contract'):
            timesheet_record.fields['print_type'] = 'Progetto N. ' + str(project_record.fields['id'])
            if project_record.fields['fixedprice'] == 'Si':
                invoicestatus = 'Fixed price Project'
                productivity = 'Ricavo indiretto'
                timesheet_record.fields['print_hourprice'] = 'Compreso nel progetto'
                timesheet_record.fields['print_travel'] = 'Inclusa'


        # valutazione flat service contract
        if invoicestatus == 'To Process':
            if not isempty(timesheet_record.fields['worktime']) and invoiceoption != 'Out of contract':
                flat_service_contract = None
                if service == 'Assistenza PBX':
                    if ((travel_time_decimal == 0 and worktime_decimal == 0.25) or invoiceoption == 'In contract'):
                        flat_service_contract = servicecontract_table.get_records(
                            conditions_list=[f"recordidcompany_='{timesheet_record.fields['recordidcompany_']}'",
                                             "(type='Manutenzione PBX')"])

                if service == 'Assistenza IT':
                    if travel_time_decimal == 0 or invoiceoption == 'In contract':
                        flat_service_contract = servicecontract_table.get_records(
                            conditions_list=[f"recordidcompany_='{timesheet_record.fields['recordidcompany_']}'",
                                             "(type='BeAll (All-inclusive)')"])

                if service == 'Printing':
                    flat_service_contract = servicecontract_table.get_records(
                        conditions_list=[f"recordidcompany_='{timesheet_record.fields['recordidcompany_']}'",
                                         "(type='Manutenzione Printing')"])

                if service == 'Assistenza Web Hosting':
                    flat_service_contract = servicecontract_table.get_records(
                        conditions_list=[f"recordidcompany_='{timesheet_record.fields['recordidcompany_']}'",
                                         "(service='Assistenza Web Hosting')"])

                if flat_service_contract:
                    servicecontract_record = Record('servicecontract', flat_service_contract[0]['recordid_'])
                    timesheet_record.fields['recordidservicecontract_'] = servicecontract_record.recordid
                    invoicestatus = 'Service Contract: ' + servicecontract_record.fields['type']
                    productivity = 'Ricavo indiretto'
                    timesheet_record.fields['print_type'] = 'Contratto di servizio'
                    timesheet_record.fields['print_hourprice'] = 'Compreso nel contratto di servizio'
                    timesheet_record.fields['print_travel'] = 'Compresa nel contratto di servizio'
        # valutazione monte ore pbx
        if ((
                invoicestatus == 'To Process' or invoicestatus == 'Under Warranty' or invoicestatus == 'Commercial support') and invoiceoption != 'Out of contract' and travel_time_decimal == 0):
            service_contracts = servicecontract_table.get_records(
                conditions_list=[f"recordidcompany_='{timesheet_record.fields['recordidcompany_']}'",
                                 "type='Monte Ore Remoto PBX'", "status='In Progress'"])
            if service_contracts:
                timesheet_record.fields['recordidservicecontract_'] = service_contracts[0]['recordid_']
                servicecontract_record = Record('servicecontract', service_contracts[0]['recordid_'])
                if invoicestatus == 'To Process':
                    invoicestatus = 'Service Contract: Monte Ore Remoto PBX'
                    productivity = 'Ricavo diretto'
                if invoicestatus == 'Under Warranty':
                    invoicestatus = 'Under Warranty'
                    productivity = 'Senza ricavo'
                if invoicestatus == 'Commercial support':
                    invoicestatus = 'Commercial support'
                    productivity = 'Senza ricavo'
                timesheet_record.fields['print_type'] = 'Monte Ore Remoto PBX'
                timesheet_record.fields['print_hourprice'] = 'Scalato dal monte ore'
                if servicecontract_record.fields['excludetravel']:
                    timesheet_record.fields['print_travel'] = 'Non scalata dal monte ore e non fatturata'
        # valutazione monte ore
        if ((
                invoicestatus == 'To Process' or invoicestatus == 'Under Warranty' or invoicestatus == 'Commercial support') and invoiceoption != 'Out of contract'):
            service_contracts = servicecontract_table.get_records(
                conditions_list=[f"recordidcompany_='{timesheet_record.fields['recordidcompany_']}'",
                                 "type='Monte Ore'", "status='In Progress'"])
            if service_contracts:
                timesheet_record.fields['recordidservicecontract_'] = service_contracts[0]['recordid_']
                servicecontract_record = Record('servicecontract', service_contracts[0]['recordid_'])
                if invoicestatus == 'To Process':
                    invoicestatus = 'Service Contract: Monte Ore'
                    productivity = 'Ricavo diretto'
                if invoicestatus == 'Under Warranty':
                    invoicestatus = 'Under Warranty'
                    productivity = 'Senza ricavo'
                if invoicestatus == 'Commercial support':
                    invoicestatus = 'Commercial support'
                    productivity = 'Senza ricavo'
                timesheet_record.fields['print_type'] = 'Monte Ore'
                timesheet_record.fields['print_hourprice'] = 'Scalato dal monte ore'
                if servicecontract_record.fields['excludetravel']:
                    timesheet_record.fields['print_travel'] = 'Non scalata dal monte ore e non fatturata'

        # da fatturare quando chiusi
        if invoicestatus == 'To Process':
            productivity = 'Ricavo diretto'
            hourprice = 140
            travelstandardprice = None
            timesheet_record.fields['print_travel'] = 'Da fatturare'

            if not isempty(company_record.fields['ictpbx_price']):
                hourprice = company_record.fields['ictpbx_price']
                travelstandardprice = company_record.fields['travel_price']

            timesheet_record.fields['print_hourprice'] = 'Fr.' + str(hourprice) + '.--'

            if not isempty(project_record.recordid):
                if project_record.fields['completed'] != 'Si':
                    invoicestatus = 'To invoice when Project Completed'

            if not isempty(ticket_record.recordid):
                if ticket_record.fields['vtestatus'] != 'Closed':
                    invoicestatus = 'To invoice when Ticket Closed'
            timesheet_record.fields['hourprice'] = hourprice
            workprice = hourprice * worktime_decimal
            timesheet_record.fields['workprice'] = workprice
            if travel_time_decimal:
                if travel_time_decimal > 0:
                    if travelstandardprice:
                        travelprice = travelstandardprice;
                    else:
                        travelprice = hourprice * travel_time_decimal;
                    timesheet_record.fields['travelprice'] = travelprice
            timesheet_record.fields['totalprice'] = workprice + travelprice
            if invoicestatus == 'To Process':
                invoicestatus = 'To Invoice'

        timesheet_record.fields['invoicestatus'] = invoicestatus
        timesheet_record.fields['productivity'] = productivity
        if service == 'Assistenza IT' or service == 'Assistenza PBX' or service == 'Assistenza SW' or service == 'Assistenza Web Hosting' or service == 'Printing':
            if timesheet_record.fields['validated'] != 'Si':
                timesheet_record.fields['validated'] = 'No'

        timesheet_record.save()

        if not isempty(servicecontract_record.recordid):
            custom_save_record(request, tableid='servicecontract', recordid=servicecontract_record.recordid)

        if not isempty(project_record.recordid):
            custom_save_record(request, tableid='project', recordid=project_record.recordid)

    # ---SERVICE CONTRACT
    if tableid == 'servicecontract':
        servicecontract_table = Table(tableid='servicecontract')
        servicecontract_record = Record('servicecontract', recordid)
        salesorderline_record = Record('salesorderline', servicecontract_record.fields['recordidsalesorderline_'])

        # recupero campi
        contracthours = servicecontract_record.fields['contracthours']
        if contracthours == None:
            contracthours = 0
        previousresidual = servicecontract_record.fields['previousresidual']
        if previousresidual == None:
            previousresidual = 0
        excludetravel = servicecontract_record.fields['excludetravel']

        # inizializzo campi
        usedhours = 0
        progress = 0
        residualhours = contracthours

        timesheet_linkedrecords = servicecontract_record.get_linkedrecords(linkedtable='timesheet')
        for timesheet_linkedrecord in timesheet_linkedrecords:
            if timesheet_linkedrecord['invoiceoption'] != 'Under Warranty' and timesheet_linkedrecord[
                'invoiceoption'] != 'Commercial support':
                usedhours = usedhours + timesheet_linkedrecord['worktime_decimal']
                if excludetravel != '1' and excludetravel != 'Si':
                    if not isempty(timesheet_linkedrecord['traveltime_decimal']):
                        usedhours = usedhours + timesheet_linkedrecord['traveltime_decimal']
        residualhours = contracthours + previousresidual - usedhours
        if contracthours + previousresidual != 0:
            progress = (usedhours / (contracthours + previousresidual)) * 100

        if isempty(servicecontract_record.fields['status']):
            servicecontract_record.fields['status'] = 'In Progress'

        servicecontract_record.fields['usedhours'] = usedhours
        servicecontract_record.fields['residualhours'] = residualhours
        servicecontract_record.fields['progress'] = progress
        servicecontract_record.save()

        if not isempty(salesorderline_record.recordid):
            custom_save_record(request, tableid='salesorderline', recordid=salesorderline_record.recordid)

    # ---SALES ORDER
    if tableid == 'salesorder':
        salesorder_record = Record('salesorder', recordid)
        deal_record = Record('deal', salesorder_record.fields['recordiddeal_'])
        salesorder_record.fields['totalcost'] = 0
        salesorder_record.fields['totalmargin'] = 0
        salesorder_record.fields['totalnetyearly'] = 0
        salesorder_record.fields['totalcostyearly'] = 0
        salesorder_record.fields['expectedmarginyearly'] = 0
        salesorder_record.fields['annual_actual_cost'] = 0
        salesorder_record.fields['totalmarginyearly'] = 0
        salesorder_record.fields['total_price'] = 0
        salesorder_record.fields['total_cost'] = 0
        salesorder_record.fields['total_margin'] = 0
        salesorder_record.fields['total_actual_cost'] = 0
        salesorder_record.fields['total_actual_margin'] = 0
        salesorder_record.fields['annual_contract_hours'] = 0
        salesorder_record.fields['annual_actual_hours'] = 0
        salesorder_record.fields['total_contract_hours'] = 0
        salesorder_record.fields['total_actual_hours'] = 0

        salesorderline_linkedrecords = salesorder_record.get_linkedrecords(linkedtable='salesorderline')
        for salesorderline_record_dict in salesorderline_linkedrecords:
            if salesorderline_record_dict['status'] == 'In Progress':
                if not isempty(salesorderline_record_dict['cost']):
                    salesorder_record.fields['totalcost'] = salesorder_record.fields['totalcost'] + \
                                                            salesorderline_record_dict['cost']
                if not isempty(salesorderline_record_dict['margin']):
                    salesorder_record.fields['totalmargin'] = salesorder_record.fields['totalmargin'] + \
                                                              salesorderline_record_dict['margin']
                if not isempty(salesorderline_record_dict['total_net_yearly']):
                    salesorder_record.fields['totalnetyearly'] = salesorder_record.fields['totalnetyearly'] + \
                                                                 salesorderline_record_dict['total_net_yearly']
                if not isempty(salesorderline_record_dict['annual_cost']):
                    salesorder_record.fields['totalcostyearly'] = salesorder_record.fields['totalcostyearly'] + \
                                                                  salesorderline_record_dict['annual_cost']
                if not isempty(salesorderline_record_dict['marginyearly']):
                    salesorder_record.fields['expectedmarginyearly'] = salesorder_record.fields[
                                                                           'expectedmarginyearly'] + \
                                                                       salesorderline_record_dict['marginyearly']
                if not isempty(salesorderline_record_dict['annual_actual_cost']):
                    salesorder_record.fields['annual_actual_cost'] = salesorder_record.fields['annual_actual_cost'] + \
                                                                     salesorderline_record_dict['annual_actual_cost']
                if not isempty(salesorderline_record_dict['annual_actual_margin']):
                    salesorder_record.fields['totalmarginyearly'] = salesorder_record.fields['totalmarginyearly'] + \
                                                                    salesorderline_record_dict['annual_actual_margin']
                if not isempty(salesorderline_record_dict['total_price']):
                    salesorder_record.fields['total_price'] = salesorder_record.fields['total_price'] + \
                                                              salesorderline_record_dict['total_price']
                if not isempty(salesorderline_record_dict['total_cost']):
                    salesorder_record.fields['total_cost'] = salesorder_record.fields['total_cost'] + \
                                                             salesorderline_record_dict['total_cost']
                if not isempty(salesorderline_record_dict['total_margin']):
                    salesorder_record.fields['total_margin'] = salesorder_record.fields['total_margin'] + \
                                                               salesorderline_record_dict['total_margin']
                if not isempty(salesorderline_record_dict['total_actual_cost']):
                    salesorder_record.fields['total_actual_cost'] = salesorder_record.fields['total_actual_cost'] + \
                                                                    salesorderline_record_dict['total_actual_cost']
                if not isempty(salesorderline_record_dict['total_actual_margin']):
                    salesorder_record.fields['total_actual_margin'] = salesorder_record.fields['total_actual_margin'] + \
                                                                      salesorderline_record_dict['total_actual_margin']
                if not isempty(salesorderline_record_dict['annual_contract_hours']):
                    salesorder_record.fields['annual_contract_hours'] = salesorder_record.fields[
                                                                            'annual_contract_hours'] + \
                                                                        salesorderline_record_dict[
                                                                            'annual_contract_hours']
                if not isempty(salesorderline_record_dict['annual_actual_hours']):
                    salesorder_record.fields['annual_actual_hours'] = salesorder_record.fields['annual_actual_hours'] + \
                                                                      salesorderline_record_dict['annual_actual_hours']
                if not isempty(salesorderline_record_dict['total_contract_hours']):
                    salesorder_record.fields['total_contract_hours'] = salesorder_record.fields[
                                                                           'total_contract_hours'] + \
                                                                       salesorderline_record_dict[
                                                                           'total_contract_hours']
                if not isempty(salesorderline_record_dict['total_actual_hours']):
                    salesorder_record.fields['total_actual_hours'] = salesorder_record.fields['total_actual_hours'] + \
                                                                     salesorderline_record_dict['total_actual_hours']

        if not isempty(deal_record.recordid):
            salesorder_record.fields['expectedmarginyearly'] = deal_record.fields['annualmargin']
        salesorder_record.save()

    # ---SALES ORDER LINE
    if tableid == 'salesorderline':
        dbh = DatabaseHelper()
        salesorderline_record = Record('salesorderline', recordid)
        salesorder_record = Record('salesorder', salesorderline_record.fields['recordidsalesorder_'])
        servicecontract_table = Table('servicecontract')
        servicecontract_record = servicecontract_table.get_record_by_condition(
            [f"recordidsalesorderline_='{recordid}'"])
        salesorderline_record.fields['repetitiontype'] = salesorder_record.fields['repetitiontype']
        salesorderline_record.fields['bexio_repetition_type'] = salesorder_record.fields['bexio_repetition_type']
        salesorderline_record.fields['bexio_repetition_interval'] = salesorder_record.fields[
            'bexio_repetition_interval']
        salesorderline_record.fields['recordidcompany_'] = salesorder_record.fields['recordidcompany_']
        salesorderline_record.fields['bexio_orderno'] = salesorder_record.fields['documentnr']
        accountid = salesorderline_record.fields['bexio_account_id']
        account_table = Table('bexio_account')
        account_record = account_table.get_record_by_condition([f"account_id='{accountid}'"])
        if account_record:
            salesorderline_record.fields['account_no'] = account_record.fields['account_no']
            salesorderline_record.fields['account'] = account_record.fields['name']
            salesorderline_record.fields['servicecontract_type'] = account_record.fields['servicecontract_type']
            salesorderline_record.fields['servicecontract_service'] = account_record.fields['servicecontract_service']
            salesorderline_record.fields['sector'] = account_record.fields['sector']

        unitcost = salesorderline_record.fields['unitcost']
        if not unitcost:
            unitcost = 0
        quantity = salesorderline_record.fields['quantity']
        if not quantity:
            quantity = 0

        linecost = unitcost * quantity
        salesorderline_record.fields['cost'] = linecost

        if salesorderline_record.fields['servicecontract_type']:
            salesorderline_record.fields['contracthours'] = (salesorderline_record.fields['price'] - linecost) / 110
            linecost = linecost + salesorderline_record.fields['contracthours'] * 60
            salesorderline_record.fields['potential_cost'] = linecost
            salesorderline_record.fields['margin'] = salesorderline_record.fields['price'] - linecost

        # calcolo annuali
        multiplier = salesorder_record.fields['multiplier']
        salesorderline_record.fields['total_net_yearly'] = salesorderline_record.fields['price'] * multiplier
        salesorderline_record.fields['annual_cost'] = linecost * multiplier
        salesorderline_record.fields['annual_actual_cost'] = salesorderline_record.fields['annual_cost']
        salesorderline_record.fields['marginyearly'] = salesorderline_record.fields['margin'] * multiplier
        salesorderline_record.fields['annual_actual_margin'] = salesorderline_record.fields['marginyearly']

        # calcolo totali
        repetitionstartdate = salesorder_record.fields['repetitionstartdate']
        lastoccurence = HelperView.get_last_occurrence(repetitionstartdate.strftime('%Y-%m-%d'))
        occurrencescount = HelperView.get_occurrences_count(repetitionstartdate.strftime('%Y-%m-%d'))
        repetitioncount = HelperView.get_repetition_count(repetitionstartdate.strftime('%Y-%m-%d'), 12 / multiplier)

        salesorderline_record.fields['total_price'] = salesorderline_record.fields['price'] * repetitioncount
        salesorderline_record.fields['total_cost'] = linecost * repetitioncount
        salesorderline_record.fields['total_margin'] = salesorderline_record.fields['total_price'] - \
                                                       salesorderline_record.fields['total_cost']

        # calcolo su contratti con ore comprese
        if salesorderline_record.fields['servicecontract_type']:

            salesorderline_record.fields['annual_contract_hours'] = salesorderline_record.fields[
                                                                        'total_net_yearly'] / 110
            salesorderline_record.fields['total_contract_hours'] = salesorderline_record.fields[
                                                                       'contracthours'] * repetitioncount
            usedhours = servicecontract_record.fields['usedhours']
            if not usedhours:
                usedhours = 0
            salesorderline_record.fields['total_actual_hours'] = usedhours

            today = datetime.datetime.now()
            twelve_months_ago = today - timedelta(days=365)  # Approximate (not considering leap years)
            result_row = dbh.sql_query_row(
                f"select sum(totaltime_decimal) as annualusedhours FROM user_timesheet WHERE recordidservicecontract_='{servicecontract_record.recordid}' and deleted_='N' AND date>='{twelve_months_ago.strftime('%Y-%m-%d')}'")
            if result_row:
                annualusedhours = result_row['annualusedhours']
                if not annualusedhours:
                    annualusedhours = 0
            salesorderline_record.fields['annual_actual_hours'] = annualusedhours
            salesorderline_record.fields['annual_actual_cost'] = annualusedhours * 60
            salesorderline_record.fields['annual_actual_margin'] = salesorderline_record.fields['total_net_yearly'] - \
                                                                   salesorderline_record.fields['annual_actual_cost']
            salesorderline_record.fields['total_actual_cost'] = usedhours * 60
            salesorderline_record.fields['total_actual_margin'] = salesorderline_record.fields['total_price'] - \
                                                                  salesorderline_record.fields['total_actual_cost']

        salesorderline_record.save()
        if not isempty(salesorder_record.recordid):
            custom_save_record(request, tableid='salesorder', recordid=salesorder_record.recordid)

    # ---PROJECT
    if tableid == 'project':
        project_record = Record('project', recordid)
        completed = project_record.fields['completed']
        deal_record = Record('deal', project_record.fields['recordiddeal_'])
        expectedhours = project_record.fields['expectedhours']
        usedhours = 0
        residualhours = 0
        fixedpricehours = 0
        servicecontracthours = 0
        bankhours = 0
        invoicedhours = 0
        timesheet_records_list = project_record.get_linkedrecords('timesheet')
        for timesheet_record_dict in timesheet_records_list:
            usedhours = usedhours + timesheet_record_dict['totaltime_decimal']
            if timesheet_record_dict['invoicestatus'] == 'Fixed Price Project':
                fixedpricehours = fixedpricehours + timesheet_record_dict['totaltime_decimal']
            if timesheet_record_dict['invoicestatus'] == 'Service Contract: Monte Ore':
                bankhours = bankhours + timesheet_record_dict['totaltime_decimal']
            if timesheet_record_dict['invoicestatus'] == 'Invoiced':
                invoicedhours = invoicedhours + timesheet_record_dict['totaltime_decimal']
        if expectedhours:
            residualhours = expectedhours - usedhours
        project_record.fields['usedhours'] = usedhours
        project_record.fields['residualhours'] = residualhours

        project_record.save()
        if not isempty(deal_record.recordid):
            deal_record.fields['usedhours'] = usedhours
            deal_record.fields['fixedpricehours'] = fixedpricehours
            deal_record.fields['servicecontracthours'] = servicecontracthours
            deal_record.fields['bankhours'] = bankhours
            deal_record.fields['invoicedhours'] = invoicedhours
            deal_record.fields['residualhours'] = residualhours
            deal_record.fields['projectcompleted'] = completed
            deal_record.save()
            custom_save_record(request, tableid='deal', recordid=deal_record.recordid)

    # ---DEALLINE
    if tableid == 'dealline':
        dealline_record = Record('dealline', recordid)
        custom_save_record(request, tableid='deal', recordid=dealline_record.fields['recordiddeal_'])

    # ---DEAL
    if tableid == 'deal':
        dbh = DatabaseHelper()
        # recupero informazioni necessarie
        deal_record = Record('deal', recordid)
        creation = deal_record.fields['creation_']
        deal_record.fields['opendate'] = creation.strftime("%Y-%m-%d")
        deal_user_recorddict = dbh.sql_query_row(f"select * from sys_user where id={deal_record.fields['dealuser1']}")
        deal_record.fields['adiuto_dealuser'] = deal_user_recorddict['adiutoid']
        deal_project_recorddict = dbh.sql_query_row(f"select * from user_project where recordiddeal_={recordid}")
        project_recordid = ''
        if deal_project_recorddict:
            project_recordid = deal_project_recorddict['recordid_']

        deal_price = deal_record.fields['amount']
        if not deal_price:
            deal_price = 0
        deal_price_sum = 0
        deal_expectedcost = deal_record.fields['expectedcost']
        if not deal_expectedcost:
            deal_expectedcost = 0
        deal_expectedcost_sum = 0
        deal_actualcost = 0
        deal_expectedhours = 0
        deal_usedhours = deal_record.fields['usedhours']
        if not deal_usedhours:
            deal_usedhours = 0
        deal_expectedmargin = 0
        deal_actualmargin = 0
        deal_annualprice = 0
        deal_annualcost = 0
        deal_annualmargin = 0

        deal_record.fields['fixedprice'] = 'No'
        dealline_records = deal_record.get_linkedrecords(linkedtable='dealline')
        for dealline_recorddict in dealline_records:
            dealline_recordid = dealline_recorddict['recordid_']
            product_recordid = dealline_recorddict['recordidproduct_']
            dealline_recordid = dealline_recorddict['recordid_']
            dealline_quantity = dealline_recorddict['quantity']
            dealline_price = dealline_recorddict['price']
            dealline_expectedcost = dealline_recorddict['expectedcost']
            dealline_expectedmargin = dealline_recorddict['expectedmargin']
            dealline_unitactualcost = dealline_recorddict['uniteffectivecost']
            if not dealline_unitactualcost:
                dealline_unitactualcost = 0
            dealline_frequency = dealline_recorddict['frequency']
            multiplier = 1
            if dealline_frequency == 'Annuale':
                multiplier = 1
            if dealline_frequency == 'Semestrale':
                multiplier = 2
            if dealline_frequency == 'Trimestrale':
                multiplier = 3
            if dealline_frequency == 'Bimestrale':
                multiplier = 6
            if dealline_frequency == 'Mensile':
                multiplier = 12
            deal_price_sum = deal_price_sum + dealline_price
            deal_expectedcost_sum = deal_expectedcost_sum + dealline_expectedcost
            dealline_record = Record('dealline', dealline_recordid)
            dealline_record.fields['recordidproject_'] = project_recordid

            dealline_actualcost = dealline_unitactualcost * dealline_quantity
            product_record = Record('product', product_recordid)
            product_fixedprice = 'No'
            if not isempty(product_record.recordid):
                product_fixedprice = product_record.fields['fixedprice']
            if not dealline_recorddict['expectedhours']:
                dealline_recorddict['expectedhours'] = 0
            deal_expectedhours = deal_expectedhours + dealline_recorddict['expectedhours']
            if product_fixedprice == 'Si':
                deal_record.fields['fixedprice'] = 'Si'
                if isempty(dealline_record.fields['expectedhours']):
                    dealline_record.fields['expectedhours'] = dealline_price / 140
                if deal_usedhours != 0:
                    dealline_record.fields['usedhours'] = deal_usedhours
                    dealline_actualcost = deal_usedhours * 60
                    deal_usedhours = 0
            if dealline_actualcost != 0:
                dealline_actualmargin = dealline_price - dealline_actualcost
            else:
                dealline_actualmargin = dealline_expectedmargin
            dealline_record.fields['effectivecost'] = dealline_actualcost
            dealline_record.fields['margin_actual'] = dealline_actualmargin

            if not isempty(dealline_frequency):
                dealline_record.fields['annualprice'] = dealline_price * multiplier
                if dealline_actualcost != 0:
                    dealline_record.fields['annualcost'] = dealline_actualcost * multiplier
                else:
                    dealline_record.fields['annualcost'] = dealline_expectedcost * multiplier
                dealline_record.fields['annualmargin'] = dealline_record.fields['annualprice'] - dealline_record.fields[
                    'annualcost']
                deal_annualprice = deal_annualprice + dealline_record.fields['annualprice']
                deal_annualcost = deal_annualcost + dealline_record.fields['annualcost']
                deal_annualmargin = deal_annualmargin + dealline_record.fields['annualmargin']
            dealline_record.save()

            deal_actualcost = deal_actualcost + dealline_actualcost
            deal_actualmargin = deal_actualmargin + dealline_actualmargin

        if(len(dealline_records) > 0): 
            deal_price = deal_price_sum
        if(len(dealline_records) > 0): 
            deal_expectedcost = deal_expectedcost_sum
        deal_expectedmargin = deal_price - deal_expectedcost
        if deal_actualcost == 0:
            deal_actualmargin = deal_expectedmargin

        deal_record.fields['amount'] = round(deal_price, 2)
        deal_record.fields['expectedcost'] = round(deal_expectedcost, 2)
        deal_record.fields['expectedmargin'] = round(deal_expectedmargin, 2)
        deal_record.fields['expectedhours'] = deal_expectedhours
        deal_record.fields['actualcost'] = deal_actualcost
        deal_record.fields['effectivemargin'] = deal_actualmargin
        deal_record.fields['margindifference'] = deal_actualmargin - deal_expectedmargin
        deal_record.fields['annualprice'] = deal_annualprice
        deal_record.fields['annualcost'] = deal_annualcost
        deal_record.fields['annualmargin'] = deal_annualmargin

        # valutazione step workflow
        deal_type = deal_record.fields['type']
        deal_record.fields['techvalidation'] = 'No'
        deal_record.fields['creditcheck'] = 'Si'
        deal_record.fields['project'] = 'Si'
        deal_record.fields['purchaseorder'] = 'Si'
        # default tech
        if deal_type == 'Printing':
            deal_record.fields['project_default_adiutotech'] = 1019
        if deal_type == 'Software' or deal_type == 'Hosting':
            deal_record.fields['project_default_adiutotech'] = 1011
        # tech validation
        if deal_type == 'ICT' or deal_type == 'PBX':
            deal_record.fields['techvalidation'] = 'Si'

        # credit check
        if deal_type == 'Rinnovo Monte ore' or deal_type == 'Riparazione Lenovo':
            deal_record.fields['creditcheck'] = 'No'
        if deal_record.fields['amount'] < 500:
            deal_record.fields['creditcheck'] = 'No'

        # progetto
        if deal_type == 'Aggiunta servizi' or deal_type == 'Materiale senza attività' or deal_type == 'Rinnovo Monte ore':
            deal_record.fields['project'] = 'No'
            deal_record.fields['project_default_adiutotech'] = 1019

        # ordine materiale
        if deal_type == 'Rinnovo Monte ore':
            deal_record.fields['purchaseorder'] = 'No'

        deal_record.save()

    # ---userlog Attività e Note
    if tableid == 'user_log':
        userlog_record = Record('user_log', recordid)
        salespush_record = Record('salespush', userlog_record.fields['recordidsalespush_'])
        reminder = userlog_record.fields['reminder']
        description = userlog_record.fields['description']
        if reminder:
            salespush_record.fields['recalldate'] = reminder.strftime("%Y-%m-%d")
        salespush_record.fields['lastupdate'] = description
        salespush_record.save()
        userlog_record.save()
    # ---CONTACT---
    if tableid == 'contact':
        contact_record = Record('contact', recordid)
        company_record = Record('company', contact_record.fields['recordidcompany_'])
        contact_record.fields['reference'] = ''

        if not isempty(contact_record.fields['name']):
            contact_record.fields['reference'] = contact_record.fields['reference'] + " " + contact_record.fields[
                'name']
        if not isempty(contact_record.fields['surname']):
            contact_record.fields['reference'] = contact_record.fields['reference'] + " " + contact_record.fields[
                'surname']
        # if not isempty(company_record.recordid):
        #   contact_record.fields['reference']=contact_record.fields['name']+" "+contact_record.fields['surname']+" - "+company_record.fields['companyname']
        contact_record.save()

    # ---STABILE
    if tableid == 'stabile':
        stabile_record = Record('stabile', recordid)
        if isempty(stabile_record.fields['titolo_stabile']):
            stabile_record.fields['titolo_stabile']=""
        riferimento=stabile_record.fields['titolo_stabile']+" "+stabile_record.fields['indirizzo']
        stabile_record.fields['riferimento']=riferimento
        stabile_record.save()

    return True


@login_required(login_url='/login/')
def pagination(request):
    if request.method == 'POST':
        page = request.POST.get('page')
        offset = 50 * page
        # "SELECT * FROM bixdata.user_invoice LIMIT 50 OFFSET %s;", [offset]

    return JsonResponse(page)


def get_test_calendar(request):
    return render(request, 'other/testCalendar.html')


@login_required(login_url='/login/')
def get_temp(request):
    return HttpResponse('test')


@login_required(login_url='/login/')
def get_settings(request):
    settings_list = get_user_setting_list(request)

    hv = HelperView(request)
    bl = SettingsBusinessLogic()
    sys_user_id = get_userid(request.user.id)
    context = dict()
    context['settings_list'] = settings_list

    query = "SELECT tableid FROM sys_user_table_order WHERE userid = '{}'".format(sys_user_id)
    with connection.cursor() as cursor:
        cursor.execute(
            query
        )
        tables = dictfetchall(cursor)

    if not tables:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT tableid FROM sys_user_table_order WHERE userid = 1"
            )
            tables = dictfetchall(cursor)

    query = f"SELECT tableid FROM sys_user_favorite_tables WHERE sys_user_id = {sys_user_id}"
    with connection.cursor() as cursor:
        cursor.execute(
            query
        )
        favorite_tables = dictfetchall(cursor)

    i = 0
    for table in tables:
        if i < len(favorite_tables) and table['tableid'] == favorite_tables[i]['tableid']:
            table['favorite'] = True
            i += 1
        else:
            table['favorite'] = False

    with connection.cursor() as cursor:
        cursor.execute("SELECT * FROM sys_table")
        all_tables = dictfetchall(cursor)

    for a, table in enumerate(all_tables):
        for b, t in enumerate(tables):
            if t['tableid'] == table['id']:
                tables[b]['description'] = table['description']

    context['tables'] = tables

    return render(request, 'other/settings.html', context)


def save_favorite_tables(request):
    fav_tables = request.POST.get('tables')
    fav_tables = json.loads(fav_tables)
    sys_user_id = get_userid(request.user.id)


    with connection.cursor() as cursor:
        cursor.execute(
            "DELETE FROM sys_user_favorite_tables where sys_user_id = %s",
            [sys_user_id]
        )

    with connection.cursor() as cursor:
        for table in fav_tables:
            cursor.execute(
                "INSERT INTO sys_user_favorite_tables(sys_user_id, tableid) VALUES (%s, %s)",
                [sys_user_id, table]
            )

    return JsonResponse({'success': True})


@login_required(login_url='/login/')
def get_under_construction(request):
    return render(request, 'other/under_construction.html')


@login_required(login_url='/login/')
def support(request):
    if request.method == 'POST':
        category = request.POST['category']
        description = request.POST['description']
        user = request.user.first_name + ' ' + request.user.last_name

        # Get the uploaded files
        images = request.FILES.getlist('images')

        # Create a directory to store the images
        directory = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'support_images')
        if not os.path.exists(directory):
            os.makedirs(directory)

        # Print the names of the images
        for image in images:
            print(image.name)

        # Do something with the saved images
        # ...

        print(category)
        print(description)
        print(user)
        print(images)

    return redirect('index')


@login_required(login_url='/login/')
def save_settings(request):
    id = None

    row = SysUser.objects.filter(bixid=request.user.id).values('id').first()

    if row:
        id = row['id']

    if request.method == 'POST':
        layout = request.POST.get('record_open_layout')
        theme = request.POST.get('theme')
        active_panel = request.POST.get('active_panel')

        with connection.cursor() as cursor:
            cursor.execute(
                "UPDATE v_sys_user_settings SET value = %s WHERE userid = %s AND setting = 'record_open_layout'",
                [layout, id]
            )
            cursor.execute(
                "UPDATE v_sys_user_settings SET value = %s WHERE userid = %s AND setting = 'theme'",
                [theme, id]
            )
            cursor.execute(
                "UPDATE v_sys_user_settings SET value = %s WHERE userid = %s AND setting = 'active_panel'",
                [active_panel, id]
            )

    return redirect('index')


@login_required(login_url='/login/')
def get_account(request):
    user = request.user
    context = {
        'user': user
    }

    userid = request.user.id

    date = datetime.datetime.now().strftime('%Y-%m-%d-%H:%M:%S')
    context['date'] = date

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT first_name, last_name FROM v_users WHERE id = {userid}"
        )
        user = dictfetchall(cursor)
        firstname= user[0]['first_name']
        lastname = user[0]['last_name']

        folder = 'bixdata_view/bixdata_app/static/images/avatars/' + firstname + lastname

        context['images'] = []
        context['user_folder'] = firstname + lastname

        if os.path.exists(folder):
            for file in os.listdir(folder):
                context['images'].append(file)

    return render(request, 'other/account.html', context)


@login_required(login_url='/login/')
def update_profile_pic(request):
    if request.method == 'POST':
        image = request.FILES.get('image')
        if image:
            # Save the uploaded image with the user's username as the filename
            filename = f"{request.user.username}.png"
            filepath = os.path.join(settings.MEDIA_ROOT, filename)
            with open(filepath, 'wb') as f:
                for chunk in image.chunks():
                    f.write(chunk)
            # Update the user's profile image URL in the database

    return redirect('index')


@login_required(login_url='/login/')
def admin_page(request):
    page = request.POST.get('page')

    rows = SysUserDashboard.objects.all().values()
    print(rows)  # Add this line to inspect the content of rows

    userids = [row.get('userid', None) for row in rows]
    dashboardids = [row.get('dashboardid', None) for row in rows]

    rows2 = SysUser.objects.filter(id__in=userids).values('firstname', 'lastname')

    names = [row['firstname'] + ' ' + row['lastname'] for row in rows2]

    with connection.cursor() as cursor4:
        cursor4.execute(
            "SELECT * FROM sys_view"
        )
        rows4 = dictfetchall(cursor4)

    with connection.cursor() as cursor5:
        cursor5.execute(
            "SELECT * FROM sys_report"
        )
        rows5 = dictfetchall(cursor5)

    tables = SysTable.objects.all().values('id')
    fields = SysField.objects.all().values('tableid', 'fieldid')

    with connection.cursor() as cursor3:
        cursor3.execute(
            "SELECT * FROM v_sys_dashboard_block"
        )
        rows3 = dictfetchall(cursor3)

        chart_names = [row['name'] for row in rows3]
        chart_dashboard_id = [row['dashboardid'] for row in rows3]

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM v_users where is_active = 1"
        )
        users = dictfetchall(cursor)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM sys_dashboard"
        )
        dashboards = dictfetchall(cursor)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM sys_user_dashboard"
        )
        user_dashboards = dictfetchall(cursor)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM sys_dashboard"
        )
        dashboards = dictfetchall(cursor)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM sys_dashboard_block"
        )
        dashboard_blocks = dictfetchall(cursor)

    for user in users:
        for dashboard in dashboards:
            user['dashboards'] = [dashboard for dashboard in dashboards]
            for dash in user['dashboards']:
                if dash['id'] in [user_dashboard['dashboardid'] for user_dashboard in user_dashboards if
                                  user_dashboard['userid'] == user['sys_user_id']]:
                    dash['visible'] = True
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT * FROM sys_dashboard"
            )
            dashboards = dictfetchall(cursor)

    context = {
        'userids': userids,
        'dashboardids': dashboardids,
        'names': names,
        'chart_names': chart_names,
        'chart_dashboard_id': chart_dashboard_id,
        'views': rows4,
        'reports': rows5,
        'tables': tables,
        'fields': fields,
        'users': users,
        'dashboards': dashboards,
        'dashboard_blocks': dashboard_blocks,
        'user_dashboards': user_dashboards
    }

    return render(request, f'admin_settings/{page}.html', {'context': context})


@login_required(login_url='/login/')
def save_chart_settings(request):
    if request.method == 'POST':
        userids = request.POST.getlist('userids[]')
        dashboardids = request.POST.getlist('dashboardids[]')
        names = request.POST.getlist('chartnames[]')
        chartdashboardids = request.POST.getlist('chartdashboardids[]')

        for i in range(len(userids)):
            user_id = userids[i]
            dashboard_id = dashboardids[i]
            SysUserDashboard.objects.filter(userid=user_id).update(dashboardid=dashboard_id)

        with connection.cursor() as cursor:
            for i in range(len(names)):
                name = names[i]
                dashboard_id = chartdashboardids[i]

                cursor.execute('UPDATE v_sys_dashboard_block SET dashboardid = %s WHERE name = %s',
                               [dashboard_id, name])

    return redirect('index')


@login_required(login_url='/login/')
def new_chart_block(request):
    if request.method == 'POST':
        name = request.POST.get('block_name')
        dashboard_id = request.POST.get('dashboard_id')
        view_id = request.POST.get('view_id')
        report_id = request.POST.get('report_id')

        with connection.cursor() as cursor:
            cursor.execute(
                """
                INSERT INTO sys_dashboard_block (dashboardid, name, userid, viewid, reportid)
                VALUES (%s, %s, %s, %s, %s)
                """,
                [dashboard_id, name, 1, view_id, report_id]
            )

    return redirect('index')


@login_required(login_url='/login/')
def get_record_path(request, tableid, recordid):
    userid = request.user.id
    if tableid == 'timesheet':
        content = get_block_timesheetinvoice(request, recordid, userid)
    elif tableid == 'task':
        content = get_block_task(request, recordid, userid)
    elif tableid == 'ticket':
        content = insert_timesheet(request, recordid, userid)
    else:
        content = get_block_record_card(request, tableid, recordid, userid)
    return index(request, content)


def insert_timesheet(request, ticketid, userid):
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT recordid_, email FROM user_ticket WHERE freshdeskid = '{ticketid}'"
        )
        rows = dictfetchall(cursor)
        recordid_ticket = rows[0]['recordid_']
        email = rows[0]['email']

    tableid = 'timesheet'
    contextfunction = 'insert'
    contextreference = tableid
    http_response = 'true'
    called_from = 'url'
    recordid = None

    table_block = get_records_table(request, 'company', None, None, email, '', 1, '', '')

    ticket_block = get_block_record_card(request, 'ticket', recordid_ticket, userid)
    # timesheet_block = get_block_record_fields(request, tableid, contextfunction, contextreference, recordid, userid, http_response, called_from)

    data = {
        'ticketid': ticketid,
        'recordid_ticket': recordid_ticket,
        'ticket_block': ticket_block,
        'table_block': table_block,
        'contextfunction': contextfunction,
        'tableid': tableid,
        'recordid': recordid,
        'userid': userid,
        'http_response': http_response,

    }

    content = render_to_string('other/insert_timesheet.html', data)

    return content


def get_block_task(request, recordid_task, userid):
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM user_task WHERE recordid_='{recordid_task}'"
        )
        rows = dictfetchall(cursor)
    if (rows):
        context = dict()
        context['task'] = rows[0]
        context['task_block'] = get_block_record_card(request, 'task', recordid_task, userid)
        content = render_to_string('other/check_task.html', context)
        return content
    else:
        return ''


def get_block_timesheetinvoice(request, recordid_timesheet, userid):
    timesheet_block = get_block_record_card(request, 'timesheet', recordid_timesheet, userid)
    company_block = ''
    project_block = ''
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM user_timesheet WHERE recordid_='{recordid_timesheet}'"
        )
        rows = dictfetchall(cursor)
    if (rows):
        recordid_company = rows[0]['recordidcompany_'];
        company_block = get_block_record_card(request, 'company', recordid_company, userid)

    context = dict()
    context['timesheet_block'] = timesheet_block
    context['company_block'] = company_block
    content = render_to_string('other/check_timesheetinvoice.html', context)
    return content;


@login_required(login_url='/login/')
def get_badge(request):
    return True


@login_required(login_url='/login/')
def get_timesheet_serviceassets(request):
    return HttpResponse('test')


@login_required(login_url='/login/')
def get_bixdata_updates(request):
    user_id = request.user.id
    return render(request, 'other/bixdata_updates.html', {'user_id': user_id})


@login_required(login_url='/login/')
def new_update(request):
    return render(request, 'other/bixdata_updates.html')


def stampa_timesheet(request):
    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')
    filename = request.POST.get('filename')
    completeUrl = request.POST.get('completeUrl')

    path = os.path.dirname(os.path.abspath(__file__))
    path = path.rsplit('views', 1)[0]
    filename_with_path = path + '\\static\\pdf\\' + filename

    uid = uuid.uuid4().hex

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=0,
    )

    today = datetime.date.today()
    d1 = today.strftime("%d/%m/%Y")

    qrcontent = str(tableid) + '_' + str(recordid)

    data = qrcontent
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    qr_name = 'qrcode' + uid + '.png'

    img.save(path + '\\static\\pdf\\' + qr_name)

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT   t.*,c.companyname,c.address,c.city,c.email, c.phonenumber, u.firstname, u.lastname FROM user_timesheet as t join user_company as c on t.recordidcompany_=c.recordid_ join sys_user as u on t.user = u.id WHERE t.recordid_='{recordid}'"
        )
        rows = dictfetchall(cursor)

        row = rows[0]

        for value in row:
            if row[value] is None:
                row[value] = ''

        row['recordid'] = recordid
        row['completeUrl'] = completeUrl + qr_name

        project_nr = ''
        ticket_nr = ''
        if row['recordidproject_'] is not None and row['recordidproject_'] != '' and row['recordidproject_'] != 'None':
            deal_nr = Helperdb.sql_query(f"SELECT iddeal FROM user_project WHERE recordid_ = {row['recordidproject_']}")
            iddeal = int(deal_nr[0]['iddeal'])
            iddeal = str(iddeal)
            project_nr = 'Progetto Nr ' + iddeal


        if row['recordidticket_'] is not None and row['recordidticket_'] != '' and row['recordidticket_'] != 'None':
            ticket_nr = Helperdb.sql_query(f"SELECT freshdeskid FROM user_ticket WHERE recordid_ = {row['recordidticket_']}")
            ticket_nr = ticket_nr[0]['freshdeskid']
            ticket_nr = 'Ticket Nr ' + ticket_nr


        if project_nr and ticket_nr:
            row['print_type'] = project_nr + ' - ' + ticket_nr
        elif project_nr:
            row['print_type'] = project_nr
        elif ticket_nr:
            row['print_type'] = ticket_nr


    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM user_timesheetline WHERE recordidtimesheet_='{recordid}'"
        )

        timesheetlines = dictfetchall(cursor)

        for line in timesheetlines:
            line['note'] = line['note'] or ''
            line['expectedquantity'] = line['expectedquantity'] or ''
            line['actualquantity'] = line['actualquantity'] or ''

    row['timesheetlines'] = timesheetlines

    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'

    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    content = render_to_string('pdf/timesheet.html', row)

    pdfkit.from_string(content, filename_with_path, configuration=config)

    try:
        with open(filename_with_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename={filename}'

        return response

    finally:
        os.remove(filename_with_path)
        os.remove(path + '\\static\\pdf\\' + qr_name)


def stampa_servicecontract(request):
    recordid = ''
    if request.method == 'POST':
        recordid = request.POST.get('recordid')
        filename = request.POST.get('filename')

        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT s.*,c.companyname,c.address,c.city,c.email FROM user_servicecontract as s join user_company as c on s.recordidcompany_=c.recordid_  WHERE s.recordid_='{recordid}'"
            )
            rows = dictfetchall(cursor)

            row = rows[0]
            row['recordid'] = recordid
            row['date'] = datetime.datetime.now().strftime("%d/%m/%Y")

        script_dir = os.path.dirname(os.path.abspath(__file__))
        wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'

        context = row
        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT t.*,u.firstname,u.lastname FROM user_timesheet as t join sys_user as u on t.user=u.id  WHERE t.recordidservicecontract_='{recordid}' AND deleted_='N' ORDER BY t.date asc"
            )
            timesheets = dictfetchall(cursor)
        timesheets_updated = list()
        for timesheet in timesheets:
            ticket_record = Record(tableid='ticket', recordid=timesheet['recordidticket_'])
            timesheet['ticket_subject'] = ''
            if not isempty(ticket_record.recordid):
                timesheet['ticket_subject'] = ticket_record.fields['subject']
            timesheets_updated.append(timesheet)
        context['timesheets'] = timesheets_updated

        config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
        content = render_to_string('pdf/servicecontract.html', row)

        filename_with_path = os.path.dirname(os.path.abspath(__file__))
        filename_with_path = filename_with_path.rsplit('views', 1)[0]
        filename_with_path = filename_with_path + '\\static\\pdf\\' + filename
        pdfkit.from_string(content, filename_with_path, configuration=config, options={"enable-local-file-access": ""})

        try:
            with open(filename_with_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/pdf")
                response['Content-Disposition'] = f'inline; filename={filename}'

            return response

        finally:
            os.remove(filename_with_path)
            
def get_stampa_gasolio_info(request):
    return HttpResponse('ok')

def stampa_gasolio(request):
    data={}
    filename='gasolio.pdf'
    
    recordid_stabile = ''
    if request.method == 'POST':
        recordid_stabile = request.POST.get('recordid')
        checkLetture=request.POST.get('checkLetture')
        meseLettura=request.POST.get('meseLettura')
        anno, mese = meseLettura.split('-')
    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    
    record_stabile=Record('stabile',recordid_stabile)
    data['stabile']=record_stabile.fields
    sql=f"""
SELECT t.recordid_,t.anno,t.mese,t.datalettura,t.lettura, i.riferimento, i.livellominimo, i.capienzacisterna
FROM user_letturagasolio t
INNER JOIN (
    SELECT recordidinformazionigasolio_, MAX(datalettura) AS max_datalettura
    FROM user_letturagasolio
    WHERE anno='{anno}' AND mese like '%{mese}%' AND deleted_='N' AND recordidstabile_ = '{recordid_stabile}'
    GROUP BY recordidinformazionigasolio_
    
) subquery
ON t.recordidinformazionigasolio_ = subquery.recordidinformazionigasolio_ 
   AND t.datalettura = subquery.max_datalettura
INNER JOIN user_informazionigasolio i
ON t.recordidinformazionigasolio_ = i.recordid_
WHERE t.recordidstabile_ = '{recordid_stabile}' AND t.deleted_ = 'N' 
        """
    ultimeletturegasolio = Helperdb.sql_query(sql)
    data['ultimeletturegasolio']=ultimeletturegasolio
    content = render_to_string('pdf/gasolio.html', data)

    filename_with_path = os.path.dirname(os.path.abspath(__file__))
    filename_with_path = filename_with_path.rsplit('views', 1)[0]
    filename_with_path = filename_with_path + '\\static\\pdf\\' + filename
    pdfkit.from_string(content, filename_with_path, configuration=config, options={"enable-local-file-access": ""})

    #return HttpResponse(content)

    try:
        with open(filename_with_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename={filename}'

            if checkLetture=='si':
                sql=f"UPDATE user_letturagasolio SET stato='Stampato' WHERE anno='{anno}' AND mese like '%{mese}%' AND recordidstabile_='{recordid_stabile}'"
                Helperdb.sql_execute(sql)
            return response
        return response

    finally:
        os.remove(filename_with_path)

def stampa_bollettini(request):
    data={}
    filename='bollettino.pdf'
    
    recordid_bollettino = ''
    if request.method == 'POST':
        recordid_bollettino = request.POST.get('recordid')
    record_bollettino = Record('bollettini',recordid_bollettino)
    recordid_stabile=record_bollettino.get_field('recordidstabile_')
    record_stabile=Record('stabile',recordid_stabile)
    recordid_dipendente=record_bollettino.get_field('recordiddipendente_')
    record_dipendente=Record('dipendente',recordid_dipendente)
    recordid_cliente=record_bollettino.get_field('recordidcliente_')
    record_cliente=Record('cliente',recordid_cliente)
    data['nome_cliente']=record_cliente.get_field('nome_cliente')
    data['riferimento']=record_stabile.get_field('riferimento')
    data['data']=record_bollettino.get_field('data')
    data['dipendente']=record_dipendente.get_field('nome')+' '+record_dipendente.get_field('cognome')
    data['informazioni']=record_bollettino.get_field('informazioni')
    data['contattatoda']=record_bollettino.get_field('contattatoda')
    data['causa']=record_bollettino.get_field('causa')
    data['interventorichiesto']=record_bollettino.get_field('interventorichiesto')
    data['id']=record_bollettino.get_field('id')  
    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    
    tipo_bollettino=record_bollettino.fields['tipo_bollettino']
    if tipo_bollettino=='Generico':
        content = render_to_string('pdf/bollettino1.html', data)
    if tipo_bollettino=='Sostituzione':
        content = render_to_string('pdf/bollettino2.html', data)
    if tipo_bollettino=='Pulizia':
        content = render_to_string('pdf/bollettino3.html', data)
    if tipo_bollettino=='Tinteggio':
        content = render_to_string('pdf/bollettino4.html', data)
    if tipo_bollettino=='Picchetto':
        content = render_to_string('pdf/bollettino5.html', data)

    filename_with_path = os.path.dirname(os.path.abspath(__file__))
    filename_with_path = filename_with_path.rsplit('views', 1)[0]
    filename_with_path = filename_with_path + '\\static\\pdf\\' + filename
    pdfkit.from_string(content, filename_with_path, configuration=config, options={"enable-local-file-access": ""})

    #return HttpResponse(content)

    try:
        with open(filename_with_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename={filename}'
            return response
        return response

    finally:
        os.remove(filename_with_path)


def stampa_servicecontract_test(request):
    recordid = request.POST.get('recordid')
    filename = request.POST.get('filename')

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT s.*,c.companyname,c.address,c.city,c.email FROM user_servicecontract as s join user_company as c on s.recordidcompany_=c.recordid_  WHERE s.recordid_='{recordid}'"
        )
        rows = dictfetchall(cursor)

        row = rows[0]
        row['recordid'] = recordid
        row['date'] = datetime.datetime.now().strftime("%d/%m/%Y")

    context = row
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT t.*,u.firstname,u.lastname FROM user_timesheet as t join sys_user as u on t.user=u.id  WHERE t.recordidservicecontract_='{recordid}'"
        )
        timesheets = dictfetchall(cursor)
    context['timesheets'] = timesheets

    return render(request, 'pdf/servicecontract.html', context)


def new_ticket_timesheet(request, ticket):
    userid = request.user.id
    content = get_block_ticket_timesheet(request, ticket, userid)
    return get_render_index(request, content)


def get_block_ticket_timesheet(request, ticket, userid):
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM user_ticket WHERE freshdeskid='{ticket}'"
        )
        rows = dictfetchall(cursor)
        if rows:
            recordid = rows[0]['recordid_']
            ticket_block = get_block_record_card(request, 'ticket', recordid, userid)

            context = dict()
            context['ticket_block'] = ticket_block
            content = render_to_string('other/check_ticket.html', context)

        return content


def rinnova_contratto(request):
    recordid = request.POST.get('recordid')
    contract_hours = request.POST.get('contract_hours')
    invoicenr = request.POST.get('invoicenr')

    post_data = {
        'recordid': recordid,
        'contracthours': contract_hours,
        'invoiceno': invoicenr,
        'startdate': datetime.datetime.now().strftime("%Y-%m-%d")
    }
    response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/rinnova_contratto", data=post_data)
    return HttpResponse('ok')


def sort_records(request):
    tableid = request.POST.get('tableid')
    order = request.POST.get('order')
    column_name = request.POST.get('order_field')

    if order == 'asc':
        icon = 'mdi mdi-arrow-down-thin'
    else:
        icon = 'mdi mdi-arrow-up-thin'

    return icon


@login_required
def export_excel(request):
    if request.method == 'POST':
        tableid = request.POST.get('tableid')
        master_tableid = request.POST.get('master_tableid')
        master_recordid = request.POST.get('master_recordid')
        searchTerm = request.POST.get('searchTerm')
        viewid = request.POST.get('viewid')
        order_field = request.POST.get('order_field')
        order = request.POST.get('order')
        currentpage = 0
        table_type = request.POST.get('tableType')

        post = {
            'tableid': tableid,
            'searchTerm': searchTerm,
            'viewid': viewid,
            'currentpage': currentpage,
            'order_field': order_field,
            'order': order,
            'master_tableid': master_tableid,
            'master_recordid': master_recordid,
            'userid': request.user.id
        }

        if table_type == 'report':
            response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/get_records_report", data=post)
        else:
            response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/get_records", data=post)

        response.raise_for_status()
        response_dict = response.json()

        csv_file = f"{tableid}-{uuid.uuid4().hex}.csv"

        csv_columns = [col['desc'] for count, col in enumerate(response_dict['columns']) if count > 2]

        records = [[remove_html_tags(field) for field in record[3:]] for record in response_dict['records']]

        with open(csv_file, 'w', newline='', encoding='utf-8-sig') as file:
            writer = csv.writer(file, delimiter=';')
            writer.writerow(csv_columns)
            writer.writerows(records)

        with open(csv_file, 'rb') as file:
            response = HttpResponse(file.read(), content_type='text/csv')
            response['Content-Disposition'] = 'attachment'
            response['filename'] = csv_file

        os.remove(csv_file)

        return response


def get_records_grouped(request):
    tableid = request.POST.get('tableid')
    table_name = 'user_' + tableid

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT c.companyname, COUNT(c.recordid_) AS NumeroTask, t.* FROM {table_name} AS t LEFT JOIN user_company AS c ON t.recordidcompany_ = c.recordid_ GROUP BY c.companyname"
        )
        rows = dictfetchall(cursor)

    context = dict()
    context['rows'] = rows

    return render(request, 'block/records/records_grouped.html', context)


# @user_passes_test(lambda u: u.is_superuser)
def send_active_task(request, requested_user=''):
    if requested_user != '':
        userid = requested_user

        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT * from v_users where id = %s", [userid]
            )
            users = dictfetchall(cursor)
    else:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT * from v_users where is_active=1"
            )
            users = dictfetchall(cursor)

    for user in users:
        sys_user_id = user['sys_user_id']
        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT user_task.*, v_users.username, user_company.companyname FROM user_task LEFT JOIN v_users ON user_task.creator = v_users.sys_user_id LEFT JOIN user_company on user_task.recordidcompany_ = user_company.recordid_ WHERE user={sys_user_id} and user_task.status!='Chiuso' and user_task.deleted_ = 'N'"
            )
            tasks = dictfetchall(cursor)

            cursor.execute(
                f"SELECT user_task.*, v_users.username, user_company.companyname FROM user_task LEFT JOIN v_users ON user_task.user = v_users.sys_user_id LEFT JOIN user_company on user_task.recordidcompany_ = user_company.recordid_ WHERE creator={sys_user_id} and user != {sys_user_id} and user_task.status!='Chiuso' and user_task.deleted_ = 'N' "
            )
            tasks_created = dictfetchall(cursor)

            current_date = datetime.date.today()
            one_week_ago = current_date - datetime.timedelta(days=7)

            # Format the date range in the YYYY-MM-DD format
            current_date_str = current_date.strftime("%Y-%m-%d")
            one_week_ago_str = one_week_ago.strftime("%Y-%m-%d")

            query = f"SELECT user_task.*, v_users.username, user_company.companyname FROM user_task \
                         LEFT JOIN v_users ON user_task.user = v_users.sys_user_id \
                         LEFT JOIN user_company ON user_task.recordidcompany_ = user_company.recordid_ \
                         WHERE creator={sys_user_id} AND user != {sys_user_id} \
                         AND closedate >= '{one_week_ago_str}' AND closedate <= '{current_date_str}' AND user_task.completed = 'Si' AND user_task.deleted_ = 'N'"

            cursor.execute(
                query
            )

            tasks_completed = dictfetchall(cursor)

            if tasks or tasks_created or tasks_completed:
                html_message = ""
                context = dict()
                for task in tasks:
                    if task['creator'] is not None:
                        task['creator'] = task['creator']
                    else:
                        task['creator'] = 'N/A'

                context['data'] = tasks
                context['data_created'] = tasks_created
                context['tasks_completed'] = tasks_completed
                html_message = render_to_string('other/send_active_task.html', context)

                subject = f"Report task - {user['first_name']} {user['last_name']}"
                email = user['email']
                send_email(emails=[email], subject=subject, html_message=html_message)

    return HttpResponse('ok')


def send_unique_active_task(request):
    requested_user = request.user.id
    send_active_task(request, requested_user)

    return True


def update_task_status(request):
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * from user_task"
        )
        tasks = dictfetchall(cursor)

        current_date = datetime.date.today()
        current_date_str = current_date.strftime("%Y-%m-%d")

        for task in tasks:
            if task['closedate'] < current_date_str:
                cursor.execute(
                    f"UPDATE user_task SET status = 'Scaduto' WHERE recordid_ = {task['recordid_']}"
                )

            elif current_date_str < task['closedate'] < current_date_str + datetime.timedelta(days=2):
                cursor.execute(
                    f"UPDATE user_task SET status = 'In scadenza' WHERE recordid_ = {task['recordid_']}"
                )
            else:
                cursor.execute(
                    f"UPDATE user_task SET status = 'Aperto' WHERE recordid_ = {task['recordid_']}"
                )

    return HttpResponse('ok')


def update_task_status(request):
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * from user_task where status!='Chiuso'"
        )
        tasks = dictfetchall(cursor)
    for task in tasks:
        post_data = {
            'tableid': 'task',
            'recordid': task['recordid_'],
            'fields': []
        }

        response = requests.post(f"{bixdata_server}bixdata/index.php/rest_controller/set_record", data=post_data)
    return HttpResponse('Eseguito')


def validate_timesheet(request):
    recordid = request.POST.get('recordid')

    with connection.cursor() as cursor:
        cursor.execute(
            "UPDATE user_timesheet SET validated = 'Si' WHERE recordid_ = %s", [recordid]
        )

    return HttpResponse('done')

def decline_timesheet(request):
    recordid = request.POST.get('recordid')
    decline_reason = request.POST.get('decline_reason')

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT user FROM user_timesheet WHERE recordid_ = {recordid}"
        )
        user = cursor.fetchone()
        user = user[0]

        cursor.execute(
            f"UPDATE user_timesheet SET validated = 'No', decline_note = '{decline_reason}' WHERE recordid_ = {recordid}"
        )

        cursor.execute(
            f"SELECT email from v_users WHERE sys_user_id = {user}"
        )
        email = cursor.fetchone()
        email = email[0]

        decliner = get_userid(request.user.id)

        cursor.execute(
            f"SELECT first_name, last_name FROM v_users WHERE sys_user_id = {decliner}"
        )
        decliner = dictfetchall(cursor)
        decliner = decliner[0]
        decliner = decliner['first_name'] + ' ' + decliner['last_name']



        message = render_to_string('other/declined_timesheet.html', {'decliner': decliner, 'decline_reason': decline_reason, 'recordid': recordid})

        send_email(
            emails=[email],
            subject='Timesheet rifiutato',
            html_message=message
        )


        return JsonResponse({'success': True})


def check_task_status(recordid):
    with connection.cursor() as cursor2:
        cursor2.execute(
            "SELECT * FROM user_task WHERE user_task.user != creator AND recordid_ = %s",
            [recordid]
        )
        task = dictfetchall(cursor2)
        if task:
            user = task[0]['user']
            company = task[0]['recordidcompany_']
        else:
            return HttpResponse('ok')

        cursor2.execute(
            "SELECT first_name, last_name   FROM v_users WHERE sys_user_id = %s",
            [user]
        )
        user = dictfetchall(cursor2)
        username = user[0]['first_name'] + ' ' + user[0]['last_name']

        cursor2.execute(
            "SELECT companyname FROM user_company WHERE recordid_ = %s",
            [company]
        )
        company = dictfetchall(cursor2)
        if company:
            company = company
        else:
            company = ''

    if task and 'status' in task[0]:
        status = task[0]['status']
        if status == 'Chiuso':
            with connection.cursor() as cursor:
                cursor.execute(
                    "SELECT description, email, username FROM v_users INNER JOIN user_task "
                    "ON v_users.sys_user_id = user_task.creator WHERE user_task.recordid_ = %s",
                    [recordid]
                )
                user = dictfetchall(cursor)

            if user:
                email = user[0]['email']
                description = user[0]['description']

                message = render_to_string('other/close_task.html',
                                           {'username': username, 'company': company, 'description': description,
                                            'recordid': recordid})
                send_email(
                    emails=[email],
                    subject='Task chiuso da ' + username,
                    html_message=message
                )

    return HttpResponse('ok')


def staff_only(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):

        if request.user.is_staff:
            return view_func(request, *args, **kwargs)
        else:
            return HttpResponse('Non sei autorizzato ad accedere a questa pagina')

    return wrapper


def check_mails():
    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * from user_task where status = 'Errore'"
        )
        tasks = dictfetchall(cursor)

    return True


def test_gridstack(request):
    return render(request, 'other/test_gridstack.html')


def new_block(request):
    blockid = request.POST.get('blockid')
    userid = request.POST.get('userid')
    size = request.POST.get('size')
    dashboardid = request.POST.get('dashboardid')
    with connection.cursor() as cursor:
        cursor.execute(
            "INSERT INTO sys_user_dashboard_block (userid, dashboard_block_id, dashboardid, size) VALUES (%s, %s, %s, %s)",
            [userid, blockid, dashboardid, size]
        )
    return JsonResponse({'success': True})


def remove_block(request):
    blockid = request.POST.get('blockid')
    size = request.POST.get('size')
    SysUserDashboardBlock.objects.filter(id=blockid, size=size).delete()

    return JsonResponse({'success': True})


def new_report(request):
    tableid = request.POST.get('tableid')
    report_name = request.POST.get('report_name')
    fieldid = request.POST.get('fieldid')
    operation = request.POST.get('operation')
    layout = request.POST.get('layout')
    groupby = request.POST.get('groupby')
    with connection.cursor() as cursor:
        cursor.execute(
            """
            INSERT INTO sys_report (userid, tableid, name, fieldid, operation, layout, groupby)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
            """,
            [1, tableid, report_name, fieldid, operation, layout, groupby]
        )

    return JsonResponse({'success': True})


def new_view(request):
    tableid = request.POST.get('tableid')
    view_name = request.POST.get('view_name')
    query_conditions = request.POST.get('query_conditions')
    with connection.cursor() as cursor:
        cursor.execute(
            """
            INSERT INTO sys_view (userid, name, tableid, query_conditions)
            VALUES (%s, %s, %s, %s)
            """,
            [1, view_name, tableid, query_conditions]
        )

    return JsonResponse({'success': True})


def test_select(request):
    return render(request, 'other/test.html')


def order_settings(request):
    tables = SysTable.objects.all().values('id')

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM v_users WHERE is_active = 1"
        )
        users = dictfetchall(cursor)

    return render(request, 'other/order_settings.html', {'tables': tables, 'users': users})


def get_table_fields(request):
    if request.method == 'POST':
        tableid = request.POST.get('tableid')
        typepreference = request.POST.get('type')
        with connection.cursor() as cursor:
            cursor.execute(

                f"SELECT sys_field.*,sys_user_order.fieldorder,sys_user_order.id AS fieldid_real FROM sys_field LEFT JOIN  (SELECT * FROM sys_user_order WHERE typepreference = '{typepreference}' ) AS sys_user_order ON sys_field.fieldid = sys_user_order.fieldid  AND sys_field.tableid = sys_user_order.tableid WHERE  sys_field.tableid = '{tableid}'"

            )
            fields = dictfetchall(cursor)

            for field in fields:
                if field['fieldid_real'] is None:
                    field['fieldid_real'] = 'None'

            returned = render_to_string('other/settings_fields.html', {'fields': fields})

            return HttpResponse(returned)


def save_fields_order(request):
    if request.method == 'POST':
        tableid = request.POST.get('tableid')
        fields_json = request.POST.get('fields')
        fields = json.loads(fields_json)

        # with connection.cursor() as cursor:
        # for count, field in enumerate(fields):
        # cursor.execute(
        #   f"UPDATE sys_user_order SET userid = 1, fieldorder = {count} WHERE tableid = '{tableid}' AND fieldid = '{field['id']}'"
        # )

        return JsonResponse({'success': True})


def update_pending_timesheet(request):
    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM user_timesheet WHERE deleted_ = 'N' AND status = 'To Invoice when Ticket Closed' AND date >= '2023-07-01'"
        )
        timesheets = dictfetchall(cursor)

        for timesheet in timesheets:
            post_data = {
                'tableid': timesheet['tableid'],
                'recordid': timesheet['recordid_'],
            }

            response = requests.post(
                f"{bixdata_server}bixdata/index.php/rest_controller/set_record", data=post_data)

    return HttpResponse('ok')


def get_project_id(request):
    if request.method == 'POST':
        projectid = request.POST.get('projectid')
        print(projectid)
        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT id, name, quantity FROM user_dealline WHERE recordidproject_ = '{projectid}'"
            )
            rows = dictfetchall(cursor)
            print(rows)

            return JsonResponse({'rows': rows})


def testtest(request):
    return render(request, 'other/test_lock.html')


class Locker:
    def __init__(self):
        self.locks = {}  # Dictionary to store lock information
        self.lock = threading.Lock()  # Create a lock for synchronization

    def acquire_lock(self, recordid, tableid, user):
        lock_key = (recordid, tableid)

        with self.lock:  # Use the lock for thread safety
            # Check if the lock is already acquired
            if lock_key in self.locks:
                return False, self.locks[lock_key]['user'], self.locks[lock_key]['timestamp']

            # If not acquired, create the lock with a timestamp
            timestamp = time.time()
            self.locks[lock_key] = {'user': user, 'timestamp': timestamp}

        return True, None, timestamp

    def release_lock(self, recordid, tableid, user):
        lock_key = (recordid, tableid)

        with self.lock:  # Use the lock for thread safety
            # Check if the lock exists and the user matches
            if lock_key in self.locks:
                if self.locks[lock_key]['user'] == user:
                    del self.locks[lock_key]
                    return True, None

        # Lock doesn't exist or user doesn't match, nothing to release
        return False, None


# Example of how to use the Locker class
locker = Locker()


def test_lock(request):
    if request.method == 'GET':

        recordid = request.GET.get('recordid')
        tableid = request.GET.get('tableid')
        userid = request.user.id  # Replace this with your actual user identification

        with lock:  # Acquire the global lock to ensure only one user at a time
            success, lock_user, timestamp = locker.acquire_lock(recordid, tableid, userid)
            print(lock)

        if success:
            return JsonResponse({'success': True})
        else:

            with connection.cursor() as cursor:
                cursor.execute(
                    f"SELECT username FROM v_users WHERE id = {lock_user}"
                )
                user_inside = cursor.fetchone()[0]

            return JsonResponse({'success': False, 'user': user_inside, 'timestamp': timestamp})

    elif request.method == 'POST':
        recordid = request.POST.get('recordid')
        tableid = request.POST.get('tableid')

        userid = request.user.id  # Replace this with your actual user identification

        with lock:  # Acquire the global lock to ensure only one user at a time
            success = locker.release_lock(recordid, tableid, userid)

        if success:
            return JsonResponse({'success': True})
        else:
            return JsonResponse({'success': False})


# Global lock for ensuring only one user can access the Locker class at a time
lock = threading.Lock()


def admin_table_settings(request):
    tables = SysTable.objects.all().values('id')
    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM v_users WHERE is_active = 1"
        )
        users = dictfetchall(cursor)

    return render(request, 'other/admin_table_settings.html', {'tables': tables, 'users': users})


def settings_charts(request):
    context = []
    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM sys_dashboard ORDER BY name asc"
        )
        rows = dictfetchall(cursor)

        with connection.cursor as cursor:
            cursor.execute(
                "SELECT * FROM v_users where is_active = 1"

            )
            users = dictfetchall(cursor)

    context['dashboards'] = rows
    context['users'] = users
    return render(request, 'admin_settings/settings_charts.html', {'context': context})


def test_adiuto_db(request):
    with connections['adiuto'].cursor() as cursor:
        cursor.execute("SELECT * FROM A1001")
        rows = dictfetchall(cursor)
    return render(request, 'other/test_adiuto_db.html')


def test_admin_doc(request):
    """Esegue una funzione

    Args:
        request (_type_): richiesta d'origin
        function (str): funzione da eseguire

    Returns:
        str: risultato della funzione chiamata, visualizzabile  nel browser
    """
    return HttpResponse('ok')


def time_calc(request):
    return render(request, 'other/time_calculator.html')


def print_word(request):
    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')
    format = request.POST.get('format')

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=0,
    )

    today = datetime.date.today()
    d1 = today.strftime("%d/%m/%Y")

    qrcontent = str(tableid) + '_' + str(recordid)

    data = qrcontent
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    qr_name = 'qrcode' + uuid.uuid4().hex + '.png'

    img.save(qr_name)

    recordid_deal = request.POST.get('recordid')
    deal_record = Record('deal', recordid_deal)
    dealuser1 = deal_record.fields['dealuser1']
    closedate = deal_record.fields['closedate']
    dealline_records = deal_record.get_linkedrecords('dealline')

    dealname = deal_record.fields['dealname']
    amount = deal_record.fields['amount']
    company_record = Record('company', deal_record.fields['recordidcompany_'])

    deal_description = deal_record.fields['description']

    companyname = company_record.fields['companyname']
    address = company_record.fields['address']
    city = company_record.fields['city']

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT first_name, last_name FROM v_users WHERE sys_user_id ='{dealuser1}'"
        )
        rows = dictfetchall(cursor)

        user = rows[0]['first_name'] + ' ' + rows[0]['last_name']

    id = uuid.uuid4().hex

    filename = dealname + id + '.docx'
    filename = filename.replace("/", "-")
    filename = filename.replace("\\", "-")
    filename = filename.replace("'", "")
    filename = filename.replace('"', "")

    # instead of creating a word i want to open one

    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Navigate to the 'views' directory and locate 'template.docx'
    file_path = os.path.join(script_dir, 'template.docx')

    # doc = Document(file_path)

    doc = Document(file_path)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    picture = run.add_picture(qr_name, width=Inches(1))

    # Set the paragraph alignment to right
    paragraph.alignment = 2  # 2 corresponds to the right alignment

    # Set spacing to minimize any additional space
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)

    os.remove(qr_name)

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)

    grey = RGBColor(0x89, 0x89, 0x89)

    p1 = doc.add_paragraph()
    text1 = f"Spett.le"
    run1 = p1.add_run(text1)
    font1 = run1.font
    font1.size = Pt(10.5)
    font1.name = 'Calibri'
    font1.bold = False
    font1.color.rgb = grey
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    p_companyname = doc.add_paragraph()
    text_companyname = f"{companyname}"
    run_companyname = p_companyname.add_run(text_companyname)
    font_companyname = run_companyname.font
    font_companyname.size = Pt(12)
    font_companyname.name = 'Calibri'
    font_companyname.bold = True
    font_companyname.color.rgb = grey
    p_companyname.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p2 = doc.add_paragraph()
    text2 = f"{address}, {city}"
    run2 = p2.add_run(text2)
    font2 = run2.font
    font2.size = Pt(10)
    font2.name = 'Calibri'
    font2.bold = False
    font2.color.rgb = grey
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)
    font_space = run_space.font
    font_space.size = Pt(10)
    font_space.name = 'Calibri'
    font_space.bold = False
    font_space.italic = True
    p_space.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_date = doc.add_paragraph()
    text_date = f"Massagno {d1}"
    run_date = p_date.add_run(text_date)
    font_date = run_date.font
    font_date.size = Pt(11)
    font_date.name = 'Calibri'
    font_date.bold = True
    font_date.color.rgb = grey
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p3 = doc.add_paragraph()
    text3 = dealname
    run3 = p3.add_run(text3)
    font3 = run3.font
    font3.size = Pt(16)
    font3.name = 'Lato'
    font3.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    font3.bold = True
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    """
    section = doc.sections[0]
    header = section.header
    watermark_path = 'background.jpg'  # Replace with your image path
    watermark = header.paragraphs[0].add_run().add_picture(watermark_path)
    watermark.alignment = WD_SECTION.DISTRIBUTE
    """

    """
    img_path = 'background.jpg'
    doc.add_picture(img_path, width=Inches(4))
    """
    doc.add_section(WD_SECTION.NEW_PAGE)

    p3 = doc.add_paragraph()
    text3 = 'Definizione Economica'
    run3 = p3.add_run(text3)
    font3 = run3.font
    font3.size = Pt(15)
    font3.name = 'Calibri'
    font3.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    font3.bold = False
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    dealline_records_length = len(dealline_records)

    table = doc.add_table(rows=dealline_records_length + 1, cols=4)  # +1 for header

    table.style = 'bixstyle'

    # Add the table header
    header_cells = ['Descrizione', 'Qt.', 'Prezzo unitario', 'Prezzo totale']
    for i, header_text in enumerate(header_cells):
        table.cell(0, i).text = header_text

    def format_chf(amount):
        if amount is None:
            return f"CHF 0.00"
        else:
            return f"CHF {amount:,.2f}".replace(",", "'")


    # Add data for each dealline
    for i, dealline in enumerate(dealline_records, start=1):
        row = table.rows[i]
        row.cells[0].text = str(dealline['name'])
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = "{:.2f}".format(dealline['quantity'])
        row.cells[2].text = format_chf(dealline['unitprice'])
        row.cells[3].text = format_chf (dealline['price'])
        row.cells[3].paragraphs[0].runs[0].font.bold = True

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    doc.add_page_break()

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    table2 = doc.add_table(rows=1, cols=1)

    table2.style = 'bixstyle'

    row_table2 = table2.rows[0]
    cell_table2 = row_table2.cells[0]
    cell_table2.text = 'Condizioni contrattuali di vendita'

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p5 = doc.add_paragraph()
    text5 = 'Contatti per Assistenza Tecnica:'
    run5 = p5.add_run(text5)
    font5 = run5.font
    font5.size = Pt(10)
    font5.name = 'Calibri'
    font5.bold = True
    font5.italic = False
    font5.color.rgb = grey
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p6 = doc.add_paragraph()  # Stile per elenco puntato con due punti
    text6 = '          •      Per tutte le richieste di assistenza: apertura ticket scrivendo all’indirizzo helpdesk@swissbix.ch  \n                  verrete ricontattati dal nostro servizio tecnico'
    run6 = p6.add_run(text6)
    font6 = run6.font
    font6.size = Pt(10)
    font6.name = 'Calibri'
    font6.bold = False
    font6.color.rgb = grey
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p7 = doc.add_paragraph()
    text7 = '          •      Orari di ufficio per supporto tecnico; dalle 9:00 alle 12:00 e dalle 14:00 alle 17:00'
    run7 = p7.add_run(text7)
    font7 = run7.font
    font7.size = Pt(10)
    font7.name = 'Calibri'
    font7.bold = False
    font7.color.rgb = grey
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p8 = doc.add_paragraph()
    text8 = 'Metodo di pagamento e fatturazione Hardware e Servizi:'
    run8 = p8.add_run(text8)
    font8 = run8.font
    font8.size = Pt(10)
    font8.name = 'Calibri'
    font8.bold = True
    font8.italic = False
    font8.color.rgb = grey
    p8.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p9 = doc.add_paragraph()
    text9 = '          •       Hardware e Consumabili: Acconto 50% all’ordine, Saldo a 20gg fine lavori'
    run9 = p9.add_run(text9)
    font9 = run9.font
    font9.size = Pt(10)
    font9.name = 'Calibri'
    font9.bold = False
    font9.color.rgb = grey
    p9.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p10 = doc.add_paragraph()
    text10 = '          •       Servizi a canone: Trimestrali anticipati a 20 giorni data fattura'
    run10 = p10.add_run(text10)
    font10 = run10.font
    font10.size = Pt(10)
    font10.name = 'Calibri'
    font10.bold = False
    font10.italic = False
    font10.color.rgb = grey
    p10.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p11 = doc.add_paragraph()
    text11 = 'Condizioni generali di vendita:'
    run11 = p11.add_run(text11)
    font11 = run11.font
    font11.size = Pt(10)
    font11.name = 'Calibri'
    font11.bold = True
    font11.color.rgb = grey
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p12 = doc.add_paragraph()
    text12 = '           •      condizioni generali di vendita sono visionabili al link: https://www.swissbix.ch/cgv.pdf'
    run12 = p12.add_run(text12)
    font12 = run12.font
    font12.size = Pt(10)
    font12.name = 'Calibri'
    font12.bold = False
    font12.color.rgb = grey
    p12.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p13 = doc.add_paragraph()
    text13 = '           •      La presente offerta comprende un servizio “chiavi in mano” al fine di \n                    garantire al cliente una totale garanzia della buona riuscita del progetto'
    run13 = p13.add_run(text13)
    font13 = run13.font
    font13.size = Pt(10)
    font13.name = 'Calibri'
    font13.bold = False
    font13.color.rgb = grey
    p13.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p14 = doc.add_paragraph()
    text14 = '           •      Offerta valida fino al ' + str(closedate) + ' o fino ad esaurimento scorte'
    run14 = p14.add_run(text14)
    font14 = run14.font
    font14.size = Pt(10)
    font14.name = 'Calibri'
    font14.bold = False
    font14.color.rgb = grey
    p14.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p15 = doc.add_paragraph()
    text15 = '           •      Swissbix SA non sarà ritenuta responsabile in caso di ritardi nella consegna del materiale \n                   dovuti a causa di forza maggiore o problemi legati ai fornitori dei prodotti o dei servizi logistici'
    run15 = p15.add_run(text15)
    font15 = run15.font
    font15.size = Pt(10)
    font15.name = 'Calibri'
    font15.bold = False
    font15.color.rgb = grey
    p15.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p16 = doc.add_paragraph()
    text16 = '           •      Sono esclusi dalla presente proposta commerciale:'
    run16 = p16.add_run(text16)
    font16 = run16.font
    font16.size = Pt(10)
    font16.name = 'Calibri'
    font16.bold = False
    font16.color.rgb = grey
    p15.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '                          o      Supporto, installazione ed eventuali uscite di fornitori esterni per gli applicativi \n                                  di terze parti utilizzati dal cliente'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '                          o      Lavori di cablaggio, lavori a muro di fissaggio e/o montaggio di ogni dispositivo, \n                                  lavori elettrici'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '                          o      Eventuali cavi, adattatori o convertitori che saranno fatturati a parte.'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '           •      I prezzi indicati sono Iva Esclusa'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p18 = doc.add_paragraph()
    text18 = 'Massagno' + ', ' + d1
    run18 = p18.add_run(text18)
    font18 = run18.font
    font18.size = Pt(10)
    font18.name = 'Calibri'
    font18.bold = False
    font18.italic = False
    font18.color.rgb = grey
    p18.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p19 = doc.add_paragraph()
    text19 = user
    run19 = p19.add_run(text19)
    font19 = run19.font
    font19.size = Pt(10)
    font19.name = 'Calibri'
    font19.bold = False
    font19.italic = False
    font19.color.rgb = grey
    p19.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p20 = doc.add_paragraph()
    text20 = 'Per Accettazione'
    run20 = p20.add_run(text20)
    font20 = run20.font
    font20.size = Pt(10)
    font20.name = 'Calibri'
    font20.bold = False
    font20.italic = False
    font20.color.rgb = grey
    p20.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p21 = doc.add_paragraph()
    text21 = '─────────────────────────────────────'
    run21 = p21.add_run(text21)
    font21 = run21.font
    font21.size = Pt(10)
    font21.name = 'Calibri'
    font21.bold = True
    font21.italic = False
    p21.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Access the section of the document (assuming there's only one section)
    section = doc.sections[0]

    # Create a footer
    footer = section.footer

    # Add a paragraph to the footer
    p22 = footer.paragraphs[0]
    text22 = 'Swissbix SA Via Baroffio 6, 6900 Lugano E-Mail: finance@swissbix.ch Telefono: +41 91 960 22 00 Banca: UBS Switzerland AG \n Titolare del conto: Swissbix SA BIC: UBSWCHZH80A IBAN: CH62 0024 7247 2096 9101 U N. IVA UE: CHE-136.887.933 '
    run22 = p22.add_run(text22)

    # Set font properties
    font22 = run22.font
    font22.size = Pt(8)
    font22.name = 'Calibri'
    font22.bold = False
    font22.italic = False

    # Set text color to grey
    font22.color.rgb = grey

    # Set paragraph alignment to center
    p22.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(filename)

    try:
        if format == 'pdf':
            with tempfile.TemporaryDirectory() as tmp_dir:
                pdf_filename = f"{tmp_dir}/{dealname}.pdf"
                docx2pdf_convert(filename, pdf_filename)

                with open(pdf_filename, 'rb') as fh:
                    response = HttpResponse(fh.read(), content_type="application/pdf")
                    response['Content-Disposition'] = f'inline; filename={dealname}.pdf'

                return response


        else:
            with open(filename, 'rb') as fh:
                response = HttpResponse(fh.read(),
                                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                response['Content-Disposition'] = f'inline; filename={dealname}.docx'

            return response

    finally:
        os.remove(filename)


def print_word_2(request):
    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')
    format = request.POST.get('format')

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=0,
    )

    today = datetime.date.today()
    d1 = today.strftime("%d/%m/%Y")

    qrcontent = str(tableid) + '_' + str(recordid)

    data = qrcontent
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    qr_name = 'qrcode' + uuid.uuid4().hex + '.png'

    img.save(qr_name)

    recordid_deal = request.POST.get('recordid')
    deal_record = Record('deal', recordid_deal)
    dealuser1 = deal_record.fields['dealuser1']
    closedate = deal_record.fields['closedate']
    dealline_records = deal_record.get_linkedrecords('dealline')

    dealname = deal_record.fields['dealname']
    amount = deal_record.fields['amount']
    company_record = Record('company', deal_record.fields['recordidcompany_'])

    deal_description = deal_record.fields['description']

    companyname = company_record.fields['companyname']
    address = company_record.fields['address']
    city = company_record.fields['city']

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT first_name, last_name FROM v_users WHERE sys_user_id ='{dealuser1}'"
        )
        rows = dictfetchall(cursor)

        user = rows[0]['first_name'] + ' ' + rows[0]['last_name']

    id = uuid.uuid4().hex

    filename = dealname + id + '.docx'
    filename = filename.replace("/", "-")
    filename = filename.replace("\\", "-")
    filename = filename.replace("'", "")

    # instead of creating a word i want to open one

    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Navigate to the 'views' directory and locate 'template.docx'
    file_path = os.path.join(script_dir, 'template2.docx')

    # doc = Document(file_path)

    doc1 = Document(file_path)

    paragraph = doc1.add_paragraph()
    run = paragraph.add_run()
    picture = run.add_picture(qr_name, width=Inches(1))

    # Set the paragraph alignment to right
    paragraph.alignment = 2  # 2 corresponds to the right alignment
    picture.alignment = 1

    # Set spacing to minimize any additional space
    paragraph.paragraph_format.space_before = Inches(0)
    paragraph.paragraph_format.space_after = Inches(0)

    os.remove(qr_name)

    section = doc1.sections[0]
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)

    grey = RGBColor(0x89, 0x89, 0x89)

    p1 = doc1.add_paragraph()
    text1 = f"Spett.le"
    run1 = p1.add_run(text1)
    font1 = run1.font
    font1.size = Pt(10.5)
    font1.name = 'Calibri'
    font1.bold = False
    font1.color.rgb = grey
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_companyname = doc1.add_paragraph()
    text_companyname = f"{companyname}"
    run_companyname = p_companyname.add_run(text_companyname)
    font_companyname = run_companyname.font
    font_companyname.size = Pt(12)
    font_companyname.name = 'Calibri'
    font_companyname.bold = True
    font_companyname.color.rgb = grey
    p_companyname.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p2 = doc1.add_paragraph()
    text2 = f"{address}, {city}"
    run2 = p2.add_run(text2)
    font2 = run2.font
    font2.size = Pt(10)
    font2.name = 'Calibri'
    font2.bold = False
    font2.color.rgb = grey
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_space = doc1.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)
    font_space = run_space.font
    font_space.size = Pt(10)
    font_space.name = 'Calibri'
    font_space.bold = False
    font_space.italic = True
    p_space.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_date = doc1.add_paragraph()
    text_date = f"Massagno {d1}"
    run_date = p_date.add_run(text_date)
    font_date = run_date.font
    font_date.size = Pt(11)
    font_date.name = 'Calibri'
    font_date.bold = True
    font_date.color.rgb = grey
    p_date.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p3 = doc1.add_paragraph()
    text3 = dealname
    run3 = p3.add_run(text3)
    font3 = run3.font
    font3.size = Pt(16)
    font3.name = 'Lato'
    font3.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    font3.bold = True
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc1.add_page_break()

    doc1.save('doc1.docx')

    filename_master = "doc1.docx"
    filename_second_docx = os.path.join(script_dir, 'template2pt2.docx')

    # filename_master is name of the file to merge the docx file into
    master = Document(filename_master)

    composer = Composer(master)
    # filename_second_docx is the name of the second docx file
    doc2 = Document(filename_second_docx)
    # append the doc2 into the master using composer.append function
    composer.append(doc2)
    # Save the combined docx with a name
    composer.save("combined.docx")

    doc = Document("combined.docx")

    """
    section = doc.sections[0]
    header = section.header
    watermark_path = 'background.jpg'  # Replace with your image path
    watermark = header.paragraphs[0].add_run().add_picture(watermark_path)
    watermark.alignment = WD_SECTION.DISTRIBUTE
    """

    """
    img_path = 'background.jpg'
    doc.add_picture(img_path, width=Inches(4))
    """
    doc.add_section(WD_SECTION.NEW_PAGE)

    p3 = doc.add_paragraph()
    text3 = 'Definizione Economica'
    run3 = p3.add_run(text3)
    font3 = run3.font
    font3.size = Pt(15)
    font3.name = 'Calibri'
    font3.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    font3.bold = False
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    dealline_records_length = len(dealline_records)

    table = doc.add_table(rows=dealline_records_length + 1, cols=4)  # +1 for header

    table.style = 'bixstyle'

    # Add the table header
    header_cells = ['Descrizione', 'Qt.', 'Prezzo unitario', 'Prezzo totale']
    for i, header_text in enumerate(header_cells):
        table.cell(0, i).text = header_text

    # Add data for each dealline
    for i, dealline in enumerate(dealline_records, start=1):
        row = table.rows[i]
        row.cells[0].text = str(dealline['name'])
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[1].text = "{:.2f}".format(dealline['quantity'])
        row.cells[2].text = "{:.2f} CHF".format(dealline['unitprice'])
        row.cells[3].text = "{:.2f} CHF".format(dealline['price'])
        row.cells[3].paragraphs[0].runs[0].font.bold = True

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    doc.add_page_break()

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    table2 = doc.add_table(rows=1, cols=1)

    table2.style = 'bixstyle'

    row_table2 = table2.rows[0]
    cell_table2 = row_table2.cells[0]
    cell_table2.text = 'Condizioni contrattuali di vendita'

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p5 = doc.add_paragraph()
    text5 = 'Contatti per Assistenza Tecnica:'
    run5 = p5.add_run(text5)
    font5 = run5.font
    font5.size = Pt(10)
    font5.name = 'Calibri'
    font5.bold = True
    font5.italic = False
    font5.color.rgb = grey
    p5.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p6 = doc.add_paragraph()  # Stile per elenco puntato con due punti
    text6 = '          •      Per tutte le richieste di assistenza: apertura ticket scrivendo all’indirizzo helpdesk@swissbix.ch  \n                  verrete ricontattati dal nostro servizio tecnico'
    run6 = p6.add_run(text6)
    font6 = run6.font
    font6.size = Pt(10)
    font6.name = 'Calibri'
    font6.bold = False
    font6.color.rgb = grey
    p6.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p7 = doc.add_paragraph()
    text7 = '          •      Orari di ufficio per supporto tecnico; dalle 9:00 alle 12:00 e dalle 14:00 alle 17:00'
    run7 = p7.add_run(text7)
    font7 = run7.font
    font7.size = Pt(10)
    font7.name = 'Calibri'
    font7.bold = False
    font7.color.rgb = grey
    p7.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p8 = doc.add_paragraph()
    text8 = 'Metodo di pagamento e fatturazione Hardware e Servizi:'
    run8 = p8.add_run(text8)
    font8 = run8.font
    font8.size = Pt(10)
    font8.name = 'Calibri'
    font8.bold = True
    font8.italic = False
    font8.color.rgb = grey
    p8.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p9 = doc.add_paragraph()
    text9 = '          •       Hardware e Consumabili: Acconto 50% all’ordine, Saldo a 20gg fine lavori'
    run9 = p9.add_run(text9)
    font9 = run9.font
    font9.size = Pt(10)
    font9.name = 'Calibri'
    font9.bold = False
    font9.color.rgb = grey
    p9.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p10 = doc.add_paragraph()
    text10 = '          •       Servizi a canone: Trimestrali anticipati a 20 giorni data fattura'
    run10 = p10.add_run(text10)
    font10 = run10.font
    font10.size = Pt(10)
    font10.name = 'Calibri'
    font10.bold = False
    font10.italic = False
    font10.color.rgb = grey
    p10.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p11 = doc.add_paragraph()
    text11 = 'Condizioni generali di vendita:'
    run11 = p11.add_run(text11)
    font11 = run11.font
    font11.size = Pt(10)
    font11.name = 'Calibri'
    font11.bold = True
    font11.color.rgb = grey
    p11.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p12 = doc.add_paragraph()
    text12 = '           •      condizioni generali di vendita sono visionabili al link: https://www.swissbix.ch/cgv.pdf'
    run12 = p12.add_run(text12)
    font12 = run12.font
    font12.size = Pt(10)
    font12.name = 'Calibri'
    font12.bold = False
    font12.color.rgb = grey
    p12.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p13 = doc.add_paragraph()
    text13 = '           •      La presente offerta comprende un servizio “chiavi in mano” al fine di \n                    garantire al cliente una totale garanzia della buona riuscita del progetto'
    run13 = p13.add_run(text13)
    font13 = run13.font
    font13.size = Pt(10)
    font13.name = 'Calibri'
    font13.bold = False
    font13.color.rgb = grey
    p13.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p14 = doc.add_paragraph()
    text14 = '           •      Offerta valida fino al ' + str(closedate) + ' o fino ad esaurimento scorte'
    run14 = p14.add_run(text14)
    font14 = run14.font
    font14.size = Pt(10)
    font14.name = 'Calibri'
    font14.bold = False
    font14.color.rgb = grey
    p14.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p15 = doc.add_paragraph()
    text15 = '           •      Swissbix SA non sarà ritenuta responsabile in caso di ritardi nella consegna del materiale \n                   dovuti a causa di forza maggiore o problemi legati ai fornitori dei prodotti o dei servizi logistici'
    run15 = p15.add_run(text15)
    font15 = run15.font
    font15.size = Pt(10)
    font15.name = 'Calibri'
    font15.bold = False
    font15.color.rgb = grey
    p15.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p16 = doc.add_paragraph()
    text16 = '           •      Sono esclusi dalla presente proposta commerciale:'
    run16 = p16.add_run(text16)
    font16 = run16.font
    font16.size = Pt(10)
    font16.name = 'Calibri'
    font16.bold = False
    font16.color.rgb = grey
    p15.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '                          o      Supporto, installazione ed eventuali uscite di fornitori esterni per gli applicativi \n                                  di terze parti utilizzati dal cliente'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '                          o      Lavori di cablaggio, lavori a muro di fissaggio e/o montaggio di ogni dispositivo, \n                                  lavori elettrici'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '                          o      Eventuali cavi, adattatori o convertitori che saranno fatturati a parte.'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p17 = doc.add_paragraph()
    text17 = '           •      I prezzi indicati sono Iva Esclusa'
    run17 = p17.add_run(text17)
    font17 = run17.font
    font17.size = Pt(10)
    font17.name = 'Calibri'
    font17.bold = False
    font17.color.rgb = grey
    p17.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p18 = doc.add_paragraph()
    text18 = 'Massagno' + ', ' + d1
    run18 = p18.add_run(text18)
    font18 = run18.font
    font18.size = Pt(10)
    font18.name = 'Calibri'
    font18.bold = False
    font18.italic = False
    font18.color.rgb = grey
    p18.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p19 = doc.add_paragraph()
    text19 = user
    run19 = p19.add_run(text19)
    font19 = run19.font
    font19.size = Pt(10)
    font19.name = 'Calibri'
    font19.bold = False
    font19.italic = False
    font19.color.rgb = grey
    p19.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p20 = doc.add_paragraph()
    text20 = 'Per Accettazione'
    run20 = p20.add_run(text20)
    font20 = run20.font
    font20.size = Pt(10)
    font20.name = 'Calibri'
    font20.bold = False
    font20.italic = False
    font20.color.rgb = grey
    p20.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    p_space = doc.add_paragraph()
    text_space = ''
    run_space = p_space.add_run(text_space)

    p21 = doc.add_paragraph()
    text21 = '─────────────────────────────────────'
    run21 = p21.add_run(text21)
    font21 = run21.font
    font21.size = Pt(10)
    font21.name = 'Calibri'
    font21.bold = True
    font21.italic = False
    p21.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Access the section of the document (assuming there's only one section)
    section = doc.sections[0]

    # Create a footer
    footer = section.footer

    # Add a paragraph to the footer
    p22 = footer.paragraphs[0]
    text22 = 'Swissbix SA Via Baroffio 6, 6900 Lugano E-Mail: finance@swissbix.ch Telefono: +41 91 960 22 00 Banca: UBS Switzerland AG \n Titolare del conto: Swissbix SA BIC: UBSWCHZH80A IBAN: CH62 0024 7247 2096 9101 U N. IVA UE: CHE-136.887.933 '
    run22 = p22.add_run(text22)

    # Set font properties
    font22 = run22.font
    font22.size = Pt(8)
    font22.name = 'Calibri'
    font22.bold = False
    font22.italic = False

    # Set text color to grey
    font22.color.rgb = grey

    # Set paragraph alignment to center
    p22.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.save(filename)

    try:
        if format == 'pdf':
            with tempfile.TemporaryDirectory() as tmp_dir:
                pdf_filename = f"{tmp_dir}/{dealname}.pdf"
                docx2pdf_convert(filename, pdf_filename)

                with open(pdf_filename, 'rb') as fh:
                    response = HttpResponse(fh.read(), content_type="application/pdf")
                    response['Content-Disposition'] = f'inline; filename={dealname}.pdf'

                return response


        else:
            with open(filename, 'rb') as fh:
                response = HttpResponse(fh.read(),
                                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                response['Content-Disposition'] = f'inline; filename={dealname}.docx'

            return response

    finally:
        os.remove(filename)


def get_record(request):
    tableid = request.POST.get('tableid')
    recordid = request.POST.get('recordid')
    r = Record(tableid, recordid)
    return_value = dict()
    if tableid == 'product':
        return_value['unitprice'] = r.fields['price']
        return_value['unitcost'] = r.fields['cost']
        return_value['name'] = r.fields['name']

    return JsonResponse(return_value)


def link_file(request):
    if request.method == 'POST' and 'file' in request.FILES:
        file = request.FILES['file']
        recordid = request.POST.get('recordid')
        tableid = request.POST.get('tableid')

        # Constructing the filename as tableid_recordid
        filename = f"{tableid}_{recordid}.{file.name.split('.')[-1]}"  # Retaining the original file extension

        # Save the file in the media root without any subdirectory
        fs = FileSystemStorage(location='deal_documents')
        filename = fs.save(filename, file)
        uploaded_file_url = fs.url(filename)

        uploaded_file_path = os.path.join('bixdata_view', filename)

        return JsonResponse({'uploaded_file_url': uploaded_file_url, 'uploaded_file_path': uploaded_file_path})
    else:
        return JsonResponse({'error': 'No file found in the request'}, status=400)


def deal_close_won(request):
    if request.method == 'POST':
        recordid = request.POST.get('recordid')
        deal_record = Record(tableid='deal', recordid=recordid)
        deal_record.fields['dealstage'] = 'Chiuso vinto'
        deal_record.fields['sync_adiuto'] = 'Si'
        deal_record.fields['dealstatus'] = 'Vinta'
        today = datetime.date.today()
        today = today.strftime("%Y-%m-%d")
        deal_record.fields['closedate'] = today
        deal_record.save()
    return JsonResponse({'success': True})


def deal_close_lost(request):
    if request.method == 'POST':
        recordid = request.POST.get('recordid')
        deal_record = Record(tableid='deal', recordid=recordid)
        deal_record.fields['dealstage'] = 'Chiuso perso'
        deal_record.fields['sync_adiuto'] = 'No'
        deal_record.fields['dealstatus'] = 'Persa'
        today = datetime.date.today()
        today = today.strftime("%Y-%m-%d")
        deal_record.fields['closedate'] = today
        deal_record.save()
    return JsonResponse({'success': True})


def deal_update_dealstage(request):
    if request.method == 'POST':
        recordid = request.POST.get('recordid')
        dealstage = request.POST.get('dealstage')
        deal_record = Record(tableid='deal', recordid=recordid)
        deal_record.fields['dealstage'] = dealstage
        deal_record.save()
    return JsonResponse({'success': True})


def custom_update(tableid, recordid):
    if (tableid == 'deal'):
        deal_table = Table(tableid='deal')
        deal_record = Record(tableid='deal', recordid=recordid)
        calc_amount = 0
        calc_expectedcost = 0
        dealline_table = Table(tableid='dealline')
        conditions_list = list()
        conditions_list.append(f"recordiddeal_='{recordid}'")
        dealline_records = dealline_table.get_records(conditions_list=conditions_list)
        for dealline_record in dealline_records:
            calc_amount = calc_amount + dealline_record['price']
            calc_expectedcost = calc_expectedcost + dealline_record['expectedcost']
        if calc_amount != 0:
            deal_record.fields['amount'] = calc_amount
        if calc_expectedcost != 0:
            deal_record.fields['expectedcost'] = calc_expectedcost
        deal_record.fields['expectedmargin'] = deal_record.fields['amount'] - deal_record.fields['expectedcost']
        deal_record.save()
    return True


def signature_function(request):
    return render(request, 'other/signature.html')


def save_signature(request):
    print('funzione signature')

    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')
    signature = request.POST.get('signature')
    completeUrl = request.POST.get('completeUrl')
    filename = request.POST.get('filename')

    print(tableid)
    print(signature)

    # download the image
    format, imgstr = signature.split(';base64,')
    ext = format.split('/')[-1]
    filename_signature = f"{tableid}_{recordid}.{ext}"

    filepath_signature = os.path.dirname(os.path.abspath(__file__))
    filepath_signature = filepath_signature.rsplit('views', 1)[0]
    filepath_signature = filepath_signature + '\\static\\pdf\\' + filename_signature

    # save the image

    with open(filepath_signature, 'wb') as fh:
        fh.write(base64.b64decode(imgstr))

    path = os.path.dirname(os.path.abspath(__file__))
    path = path.rsplit('views', 1)[0]
    filename_with_path = path + '\\static\\pdf\\' + filename

    uid = uuid.uuid4().hex

    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_H,
        box_size=10,
        border=0,
    )

    today = datetime.date.today()
    d1 = today.strftime("%d/%m/%Y")

    qrcontent = str(tableid) + '_' + str(recordid)

    data = qrcontent
    qr.add_data(data)
    qr.make(fit=True)

    img = qr.make_image(fill_color="black", back_color="white")

    qr_name = 'qrcode' + uid + '.png'

    img.save(path + '\\static\\pdf\\' + qr_name)

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT   t.*,c.companyname,c.address,c.city,c.email, c.phonenumber, u.firstname, u.lastname FROM user_timesheet as t join user_company as c on t.recordidcompany_=c.recordid_ join sys_user as u on t.user = u.id WHERE t.recordid_='{recordid}'"
        )
        rows = dictfetchall(cursor)

        row = rows[0]

        for value in row:
            if row[value] is None:
                row[value] = ''

        row['recordid'] = recordid
        row['completeQrUrl'] = completeUrl + qr_name
        row['completeSignatureUrl'] = completeUrl + filename_signature

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM user_timesheetline WHERE recordidtimesheet_='{recordid}'"
        )

        timesheetlines = dictfetchall(cursor)

        for line in timesheetlines:
            line['note'] = line['note'] or ''
            line['expectedquantity'] = line['expectedquantity'] or ''
            line['actualquantity'] = line['actualquantity'] or ''

    row['timesheetlines'] = timesheetlines

    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'

    config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
    content = render_to_string('pdf/timesheet_signature.html', row)

    pdfkit.from_string(content, filename_with_path, configuration=config)

    try:
        with open(filename_with_path, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/pdf")
            response['Content-Disposition'] = f'inline; filename={filename}'

        return response

    finally:
        os.remove(path + '\\static\\pdf\\' + qr_name)
        os.remove(filepath_signature)
        os.remove(filename_with_path)


def new_dashboard(request):
    dashboard_name = request.POST.get('dashboard-name')
    with connection.cursor() as cursor:
        cursor.execute(
            "INSERT INTO sys_dashboard (name,userid) VALUES (%s, %s)",
            [dashboard_name, 1]
        )
    return JsonResponse({'success': True})


def save_users_dashboards(request):
    with connection.cursor() as cursor:
        cursor.execute(
            "DELETE FROM sys_user_dashboard"
        )

    for key, value in request.POST.items():
        if '-' in key:
            dashboard_name, sys_user_id = key.split('-')
            dashboard_id = value

            with connection.cursor() as cursor:
                cursor.execute(
                    "INSERT INTO sys_user_dashboard (userid,dashboardid) VALUES (%s, %s)",
                    [sys_user_id, dashboard_id]
                )

    return JsonResponse({'success': True})


def set_default_dashboard(request):
    dashboardid = request.POST.get('dashboardid')

    sys_user = get_userid(request.user.id)

    with connection.cursor() as cursor:
        cursor.execute(
            "UPDATE sys_user_settings SET value = %s WHERE setting = 'default_dashboard' AND userid = %s",
            [dashboardid, sys_user]
        )

        if cursor.rowcount == 0:
            cursor.execute(
                "INSERT INTO sys_user_settings (userid, setting, value) VALUES (%s, 'default_dashboard', %s)",
                [sys_user, dashboardid]
            )

    return JsonResponse({'success': True})


def get_company_card(request, phonenumber):
    userid = request.user.id

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM user_company WHERE phonenumber='{phonenumber}'"
        )
        rows = dictfetchall(cursor)
        rows = rows[0]

    if (rows):
        context = dict()
        context['company'] = rows
        context['company_block'] = get_block_record_card(request, 'company', rows['recordid_'], userid)
        content = render_to_string('other/company_card.html', context)
        return index(request, content)


def get_3cx_card(request, phonenumber, callername):
    userid = request.user.id
    context = dict()

    contact = Helperdb.sql_query(f"SELECT recordidcompany_ FROM user_contact WHERE phone = '{phonenumber}' OR mobilephone = '{phonenumber}'")
    if contact:
        contact = contact[0]
        rows = Helperdb.sql_query(f"SELECT * FROM user_company WHERE recordid_ = '{contact['recordidcompany_']}'")
    else:
        with connection.cursor() as cursor:
            cursor.execute(
                f"SELECT * FROM user_company WHERE phonenumber='{phonenumber}'"
            )
            rows = dictfetchall(cursor)

    if (rows):
        rows = rows[0]
        context['company'] = rows
        context['company_block'] = get_block_record_card(request, 'company', rows['recordid_'], userid)
        content = render_to_string('other/company_card.html', context)
    else:
        companies = Helperdb.sql_query("SELECT * FROM user_company")
        context['companies'] = companies
        context['phonenumber'] = phonenumber
        content = render_to_string('other/add_phonenumber.html', context)

    return index(request, content)


def save_phonenumber(request):
    recordid = request.POST.get('recordid')
    phonenumber = request.POST.get('phonenumber')

    Helperdb.sql_execute(f"UPDATE user_company SET phonenumber = '{phonenumber}' WHERE recordid_ = '{recordid}'")

    return JsonResponse({'success': True})


def notify_error(request):
    return JsonResponse({'success': True})


def save_dashboard_table(request):
    rows = request.POST.get('rows')
    rows = json.loads(rows)

    for row in rows:
        names = []
        values = []

        for value in row['values']:
            names.append(value['name'])
            values.append(value['value'])
            formatted_values = ["'{}'".format(value) for value in values]

        if row['action'] == 'add':
            query = "INSERT INTO sys_dashboard ({}) VALUES ({})".format(
                ', '.join(names), ', '.join(formatted_values)
            )
            with connection.cursor() as cursor:
                cursor.execute(query)

        if row['action'] == 'delete':
            query = "DELETE FROM sys_dashboard where id = '{}'".format(values[0])
            with connection.cursor() as cursor:
                cursor.execute(query)

        if row['action'] == 'edit':
            query = "REPLACE INTO sys_dashboard ({}) VALUES ({})".format(
                ', '.join(names), ', '.join(formatted_values)
            )
            with connection.cursor() as cursor:
                cursor.execute(query)

    return JsonResponse({'success': True})


def save_columns_width(request):
    tableid = request.POST.get('tableid')
    width = request.POST.get('widths')
    userid = get_userid(request.user.id)

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM sys_user_column_width WHERE userid = %s AND tableid = %s",
            [userid, tableid]
        )
        existing_row = cursor.fetchone()

        if existing_row:
            cursor.execute(
                "UPDATE sys_user_column_width SET column_width = %s WHERE userid = %s AND tableid = %s",
                [width, userid, tableid]
            )
        else:
            cursor.execute(
                "INSERT INTO sys_user_column_width (userid, tableid, column_width) VALUES (%s, %s, %s)",
                [userid, tableid, width]
            )

    return JsonResponse({'success': True})


def get_user_worktime(request):
    userid = get_userid(request.user.id)

    current_date = datetime.date.today().strftime('%Y-%m-%d')

    query = f"SELECT SUM(totaltime_decimal) FROM user_timesheet WHERE user = '{userid}' AND date = '{current_date}' AND deleted_ = 'N'"
    print(query)
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT SUM(totaltime_decimal) FROM user_timesheet WHERE user = '{userid}' AND date = '{current_date}' AND deleted_ = 'N'"
        )
        worked = cursor.fetchone()[0]

        if worked is None:
            worked = 0

    return JsonResponse({'worktime': worked})


def update_record_date(request):
    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')
    datetype = request.POST.get('datetype')
    date = request.POST.get('date')
    endhour = request.POST.get('endhour')
    starthour = request.POST.get('starthour')
    allDay = request.POST.get('allDay')

    if datetype == 'planneddate' and endhour and starthour:

        with connection.cursor() as cursor:
            cursor.execute(
                f"UPDATE user_task SET {datetype} = '{date}', start = '{starthour}', end = '{endhour}' WHERE recordid_ = '{recordid}'"

            )

    else:
        with connection.cursor() as cursor:
            cursor.execute(
                f"UPDATE user_task SET {datetype} = '{date}' WHERE recordid_ = '{recordid}'"
            )

    return JsonResponse({'success': True})


def update_end_date(request):
    enddate = request.POST.get('enddate')
    recordid = request.POST.get('recordid')

    with connection.cursor() as cursor:
        cursor.execute(
            f"UPDATE user_task SET end = '{enddate}' WHERE recordid_ = '{recordid}'"
        )

    print(enddate)
    print(recordid)

    return JsonResponse({'success': True})


def set_event_allday(request):
    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')
    datetype = request.POST.get('datetype')
    date = request.POST.get('date')

    with connection.cursor() as cursor:
        cursor.execute(
            f"UPDATE user_task SET {datetype} = '{date}',  start = NULL, end = NULL WHERE recordid_ = '{recordid}'"
        )

    return JsonResponse({'success': True})


def get_script(request):
    sql = "SELECT * FROM sys_table"
    with connection.cursor() as cursor:
        cursor.execute(sql)
        tables = dictfetchall(cursor)

    context = {}
    context['tables'] = tables

    return render(request, 'other/script_page.html', context)


def execute_script(request):
    tableid = request.POST.get('tableid')
    condition = request.POST.get('condition')

    recordids = []

    sql = f"SELECT * FROM user_{tableid} WHERE {condition}"
    with connection.cursor() as cursor:
        cursor.execute(sql)
        records = dictfetchall(cursor)

    for record in records:
        recordids.append(record['recordid_'])

    for recordid in recordids:
        save_record_fields(request)

    return True


def close_salespush(request):
    recordid = request.POST.get('recordid')

    with connection.cursor() as cursor:
        cursor.execute(
            f"UPDATE user_salespush SET status = 'Chiuso' WHERE recordid_ = '{recordid}'"
        )

    return JsonResponse({'success': True})


def get_user_sold(request):
    userid = get_userid(request.user.id)

    query = f"SELECT ROUND(SUM(d.effectivemargin)) FROM user_deal AS d WHERE d.dealuser1 = '{userid}' AND d.closedate >= DATE_FORMAT(CURDATE(), '%Y-%m-01') AND d.closedate <= CURDATE() AND d.deleted_ = 'N'"

    with connection.cursor() as cursor:
        cursor.execute(query)
        sold = cursor.fetchone()[0]

        if sold is None:
            sold = 0

    return JsonResponse({'sold': sold})


def task_functions(request):
    func = request.POST.get('func')
    recordid = request.POST.get('recordid')
    tableid = request.POST.get('tableid')

    if func == 'plannedtoday':
        with connection.cursor() as cursor:
            cursor.execute(
                f"UPDATE user_task SET planneddate = CURDATE() WHERE recordid_ = '{recordid}'"
            )

    elif func == 'plannedtomorrow':
        with connection.cursor() as cursor:
            cursor.execute(
                f"UPDATE user_task SET planneddate = DATE_ADD(CURDATE(), INTERVAL 1 DAY) WHERE recordid_ = '{recordid}'"
            )

    elif func == 'weekpostpone':
        with connection.cursor() as cursor:
            cursor.execute(
                f"UPDATE user_task SET planneddate = NULL, duedate = CURDATE() + INTERVAL 7 DAY WHERE recordid_ = '{recordid}'"
            )
    elif func == 'monthpostpone':
        with connection.cursor() as cursor:
            cursor.execute(
                f"UPDATE user_task SET planneddate = NULL, duedate = CURDATE() + INTERVAL 1 MONTH WHERE recordid_ = '{recordid}'"
            )

    return JsonResponse({'success': True})


def stampa_project(request):
    dbh = DatabaseHelper()
    recordid = request.POST.get('recordid')
    type = request.POST.get('type')
    users = dict()
    rows = dbh.sql_query("SELECT * FROM sys_user WHERE firstname is not null and lastname is not null")
    for row in rows:
        users[int(row['id'])] = row['firstname'] + " " + row['lastname']

    path = os.path.dirname(os.path.abspath(__file__))
    path = path.rsplit('views', 1)[0]
    filename_with_path = path + '\\static\\pdf\\' + 'project.pdf'

    uid = uuid.uuid4().hex

    context = {}

    context['type'] = type

    record_project = Record('project', recordid)
    record_company = Record('company', record_project.fields['recordidcompany_'])
    context['nomeprogetto'] = record_project.fields['projectname']
    context['nomeazienda'] = record_company.fields['companyname']
    context['responsabile'] = users[int(record_project.fields['assignedto'])]

    technicians = list()
    sql = f"""
            SELECT distinct user,COUNT(*) AS counter 
            FROM user_timesheet WHERE deleted_='N' AND recordidproject_='{recordid}' 
            GROUP BY USER
            ORDER BY counter desc
        """
    rows = dbh.sql_query(sql)
    for row in rows:
        technicians.append(users[int(row['user'])])
    context['tecnici'] = technicians
    context['datainizio'] = record_project.fields['startdate']
    context['datafineprevista'] = record_project.fields['targetenddate']
    context['datafineeffettiva'] = ''
    context['descrizione'] = null_check(record_project.fields['publicdescription']).replace('\n', '<br/>')
    context['note'] = null_check(record_project.fields['publicnote']).replace('\n', '<br/>')
    table_timesheet = Table('timesheet')
    conditions_list = list()
    conditions_list.append(f"recordidproject_ = {recordid}")
    conditions_list.append(f"deleted_='N'")
    timesheets = table_timesheet.get_records(conditions_list=conditions_list, orderby='date asc')
    context['timesheets'] = list()
    for timesheet in timesheets:
        timesheet['user'] = users[int(timesheet['user'])]
        context['timesheets'].append(timesheet)

        first_name, last_name = timesheet['user'].split(' ', 1)
        profile_pic = f"{first_name}.{last_name}"
        profile_pic = profile_pic.lower()
        timesheet['profile_pic'] = profile_pic

    table_projectmilestone = Table('projectmilestone')
    conditions_list = list()
    conditions_list.append(f"recordidproject_ = {recordid}")
    conditions_list.append(f"deleted_='N'")
    milestones = table_projectmilestone.get_records(conditions_list=conditions_list, orderby='expecteddate asc')
    context['milestones'] = milestones

    milestones_tasks = dict()
    table_tasks = Table('task')
    for milestone in milestones:
        milestone_recordid = milestone['recordid_']
        milestones_tasks[milestone_recordid] = dict()
        milestones_tasks[milestone_recordid]['titolo'] = milestone['title']
        milestones_tasks[milestone_recordid]['descrizione'] = milestone['note']
        milestones_tasks[milestone_recordid]['stato'] = milestone['stato']
        conditions_list = list()
        conditions_list.append(f"recordidprojectmilestone_ = {milestone['recordid_']}")
        milestones_tasks[milestone_recordid]['tasks'] = table_tasks.get_records(conditions_list=conditions_list,
                                                                                orderby='recordid_ asc')
    context['milestones_tasks'] = milestones_tasks
    script_dir = os.path.dirname(os.path.abspath(__file__))
    wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'

    record_deal = Record('deal', record_project.fields['recordiddeal_'])
    if not isempty(record_deal.recordid):
        context['redditivita'] = dict()
        context['redditivita']['labels'] = ['Importo trattativa', 'Margine Previsto', 'Margine effettivo']
        context['redditivita']['values'] = [record_deal.fields['amount'], record_deal.fields['expectedmargin'],
                                            record_deal.fields['effectivemargin']]
        context['redditivita']['importo'] = record_deal.fields['amount']
        context['redditivita']['margineprevisto'] = record_deal.fields['expectedmargin']
        context['redditivita']['margineffettivo'] = record_deal.fields['effectivemargin']
    else:
        context['redditivita'] = dict()
        context['redditivita']['labels']=[]
        context['redditivita']['values']=[]
        context['redditivita']['importo']=""
        context['redditivita']['margineprevisto']=""
        context['redditivita']['margineffettivo']=""
    righedettaglio = list()
    table_dealline = Table('dealline')
    conditions_list = list()
    conditions_list.append(f"recordidproject_ = {recordid}")
    conditions_list.append(f"deleted_='N'")
    righedettaglio = table_dealline.get_records(conditions_list=conditions_list, orderby='recordid_ desc')
    context['righedettaglio'] = righedettaglio

    sql = f"SELECT SUM(totaltime_decimal) as totaltime_decimal,service FROM user_timesheet WHERE deleted_='N' AND recordidproject_='{recordid}' GROUP BY service"
    rows = dbh.sql_query(sql)
    values = list()
    labels = list()
    for row in rows:
        values.append(row['totaltime_decimal'])
        labels.append(row['service'])
    if None in labels:
        labels = ['Non assegnato' if v is None else v for v in labels]
    context['ore_servizio'] = dict()
    context['ore_servizio']['labels'] = labels
    context['ore_servizio']['values'] = values

    sql = f"SELECT SUM(totaltime_decimal) as totaltime_decimal,invoiceoption FROM user_timesheet WHERE deleted_='N' AND recordidproject_='{recordid}' GROUP BY invoiceoption"
    rows = dbh.sql_query(sql)
    values = list()
    labels = list()
    for row in rows:
        values.append(row['totaltime_decimal'])
        labels.append(row['invoiceoption'])
    if None in labels:
        labels = ['Non assegnato' if v is None else v for v in labels]
    context['ore_opzioni'] = dict()
    context['ore_opzioni']['labels'] = labels
    context['ore_opzioni']['values'] = values

    """
        try:
            with open(filename_with_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/pdf")
                response['Content-Disposition'] = f'inline; filename=project.pdf'

            return response

        finally:
            os.remove(filename_with_path)
        """
    return render(request, 'pdf/project.html', context)


def get_ticket_feedback(request):
    url = f"https://www.swissbix.ch/ticketfeedback/get_tickets.php?password={get_tickets_password}"

    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'
    }

    response = requests.get(url, headers=headers)

    response = json.loads(response.text)
    for feedback in response:
        field = Helperdb.sql_query_row(f"select * from user_ticketfeedback WHERE ticketid='{feedback['ticketid']}'")
        if not field:
            new_record = Record(tableid='ticketfeedback')
            new_record.fields['ticketid'] = feedback['ticketid']
            new_record.fields['level'] = feedback['level']
            new_record.fields['comment'] = feedback['comment']

            new_record.save()
        else:
            record = Record(tableid='ticketfeedback', recordid=field['recordid_'])
            record.fields['level'] = feedback['level']
            record.fields['comment'] = feedback['comment']
            record.save()

    return JsonResponse(response, safe=False)


def get_freshdesk_tickets(request):
    api_key = os.environ.get('FRESHDESK_APIKEY')
    password = "x"
    yourdomain = "swissbix"

    url = f"https://{yourdomain}.freshdesk.com/api/v2/tickets?include=requester,description,stats&updated_since=2020-07-01&per_page=10"

    response = requests.get(url, auth=(api_key, password))

    headers = response.headers
    response = json.loads(response.text)

    for ticket in response:
        field = Helperdb.sql_query_row(f"select * from user_freshdesk_tickets WHERE ticket_id='{ticket['id']}'")
        if not field:
            new_record = Record(tableid='freshdesk_tickets')
            new_record.fields['ticket_id'] = ticket['id']
            new_record.fields['subject'] = ticket['subject']
            new_record.fields['description'] = ticket['description_text']
            new_record.fields['created_at'] = ticket['created_at']
            new_record.fields['closed_at'] = ticket['stats']['closed_at']
            new_record.fields['requester_id'] = ticket['requester']['id']
            new_record.fields['requester_name'] = ticket['requester']['name']
            new_record.fields['requester_email'] = ticket['requester']['email']
            new_record.fields['responder_id'] = ticket['responder_id']
            new_record.fields['status'] = ticket['status']

            new_record.save()
        else:
            record = Record(tableid='freshdesk_tickets', recordid=field['recordid_'])
            record.fields['subject'] = ticket['subject']
            record.fields['description'] = ticket['description_text']
            record.fields['closed_at'] = ticket['stats']['closed_at']
            record.fields['status'] = ticket['status']
            record.fields['responder_id'] = ticket['responder_id']

            record.save()


    return JsonResponse(response, safe=False)


def get_bexio_contacts(request):
    url = "https://api.bexio.com/2.0/contact/"
    accesstoken=os.environ.get('BEXIO_ACCESSTOKEN')
    headers = {
        'Accept': "application/json",
        'Content-Type': "application/json",
        'Authorization': f"Bearer {accesstoken}",
    }

    response = requests.request("GET", url, headers=headers)
    response = json.loads(response.text)

    for contact in response:
        field = Helperdb.sql_query_row(f"select * from user_bexio_contact WHERE bexio_id='{contact['id']}'")
        if not field:
            record = Record(tableid="bexio_contact")

        else:
            record = Record(tableid="bexio_contact", recordid=field['recordid_'])

        record.fields['bexio_id'] = contact['id']
        record.fields['nr'] = contact['nr']
        record.fields['nr'] = contact['nr']
        record.fields['contact_type_id'] = contact['contact_type_id']
        record.fields['name_1'] = contact['name_1']
        record.fields['name_2'] = contact['name_2']
        record.fields['address'] = contact['address']
        record.fields['postcode'] = contact['postcode']
        record.fields['city'] = contact['city']
        record.fields['country_id'] = contact['country_id']
        record.fields['mail'] = contact['mail']
        record.fields['mail_second'] = contact['mail_second']
        record.fields['phone_fixed'] = contact['phone_fixed']
        record.fields['phone_mobile'] = contact['phone_mobile']
        record.fields['contact_group_ids'] = contact['contact_group_ids']
        record.fields['contact_branch_ids'] = contact['contact_branch_ids']
        record.fields['user_id'] = contact['user_id']
        record.fields['owner_id'] = contact['owner_id']
        # record.fields['status'] = contact['status']


        record.save()

    return JsonResponse(response, safe=False)

def get_bexio_orders(request):
    url = "https://api.bexio.com/2.0/kb_order/search/?order_by=id_desc&limit=100&offset=0"
    accesstoken=os.environ.get('BEXIO_ACCESSTOKEN')
    headers = {
        'Accept': "application/json",
        'Content-Type': "application/json",
        'Authorization': f"Bearer {accesstoken}",
    }

    payload = """
    [
        {
            "field": "kb_item_status_id",
            "value": "5",
            "criteria": "="
        }
    ]
    """

    response = requests.request("POST", url, data=payload, headers=headers)
    response = json.loads(response.text)

    for order in response:
        field = Helperdb.sql_query_row(f"select * from user_bexio_orders WHERE bexio_id='{order['id']}'")
        if not field:
            record = Record(tableid="bexio_orders")

        else:
            record = Record(tableid="bexio_orders", recordid=field['recordid_'])


        record.fields['bexio_id'] = order['id']
        record.fields['document_nr'] = order['document_nr']
        record.fields['document_nr'] = order['document_nr']
        record.fields['title'] = order['title']
        record.fields['contact_id'] = order['contact_id']
        record.fields['user_id'] = order['user_id']
        record.fields['total_gross'] = order['total_gross']
        record.fields['total_net'] = order['total_net']
        record.fields['total_taxes'] = order['total_taxes']
        record.fields['total'] = order['total']
        record.fields['is_valid_from'] = order['is_valid_from']
        record.fields['contact_address'] = order['contact_address']
        record.fields['delivery_address'] = order['delivery_address']
        record.fields['is_recurring'] = order['is_recurring']
        record.fields['is_recurring'] = order['is_recurring']
        # record.fields['taxs_percentage'] = order['taxs']['percentage']
        # record.fields['taxs_value'] = order['taxs']['value']

        record.save()

        get_bexio_positions(request, 'kb_order', order['id'])

    return JsonResponse(response, safe=False)


def get_bexio_positions_example(request, bexioid):
    return get_bexio_positions(request,'kb_order',bexioid)


def get_bexio_positions(request,bexiotable,bexioid):
    url = f"https://api.bexio.com/2.0/{bexiotable}/{bexioid}/kb_position_custom"
    accesstoken=os.environ.get('BEXIO_ACCESSTOKEN')
    headers = {
        'Accept': "application/json",
        'Content-Type': "application/json",
        'Authorization': f"Bearer {accesstoken}",
    }

    response = requests.request("GET", url, headers=headers)
    response = json.loads(response.text)

    for position in response:
        field = Helperdb.sql_query_row(f"select * from user_bexio_positions WHERE bexio_id='{position['id']}'")
        if not field:
            record = Record(tableid="bexio_positions")
        else:
            record = Record(tableid="bexio_positions", recordid=field['recordid_'])


        if bexiotable == 'kb_order':
            type='order'
        else:
            type='invoice'

        record.fields['bexio_id'] = position['id']
        record.fields['type'] = type
        record.fields['amount'] = position['amount']
        record.fields['unit_id'] = position['unit_id']
        record.fields['account_id'] = position['account_id']
        record.fields['unit_name'] = position['unit_name']
        record.fields['tax_id'] = position['tax_id']
        record.fields['tax_value'] = position['tax_value']
        record.fields['text'] = position['text']
        record.fields['unit_price'] = position['unit_price']
        record.fields['discount_in_percent'] = position['discount_in_percent']
        record.fields['position_total'] = position['position_total']
        record.fields['pos'] = position['pos']
        record.fields['internal_pos'] = position['internal_pos']
        record.fields['parent_id'] = position['parent_id']
        record.fields['is_optional'] = position['is_optional']

        record.save()

    return JsonResponse(response, safe=False)


def get_bexio_invoices(request):
    url = "https://api.bexio.com/2.0/kb_invoice?order_by=id_desc&limit=100&offset=0"
    accesstoken=os.environ.get('BEXIO_ACCESSTOKEN')
    headers = {
        'Accept': "application/json",
        'Content-Type': "application/json",
        'Authorization': f"Bearer {accesstoken}",
    }

    response = requests.request("GET", url, headers=headers)
    response = json.loads(response.text)

    for invoice in response:
        field = Helperdb.sql_query_row(f"select * from user_bexio_invoices WHERE bexio_id='{invoice['id']}'")
        if not field:
            record = Record(tableid="bexio_invoices")
        else:
            record = Record(tableid="bexio_invoices", recordid=field['recordid_'])

        record.fields['bexio_id'] = invoice['id']
        record.fields['document_nr'] = invoice['document_nr']
        record.fields['title'] = invoice['title']
        record.fields['contact_id'] = invoice['contact_id']
        record.fields['user_id'] = invoice['user_id']
        record.fields['total_gross'] = invoice['total_gross']
        record.fields['total_net'] = invoice['total_net']
        record.fields['total_taxes'] = invoice['total_taxes']
        record.fields['total_received_payments'] = invoice['total_received_payments']
        record.fields['total_remaining_payments'] = invoice['total_remaining_payments']
        record.fields['total'] = invoice['total']
        record.fields['is_valid_from'] = invoice['is_valid_from']
        record.fields['is_valid_to'] = invoice['is_valid_to']
        record.fields['contact_address'] = invoice['contact_address']


        record.save()

        get_bexio_positions(request, 'kb_invoice', invoice['id'])


    return JsonResponse(response, safe=False)

def syncdata(request,tableid):
    sync_table=Helperdb.db_get_value('sys_table','sync_table',f"id='{tableid}'")
    sync_field=Helperdb.db_get_value('sys_table','sync_field',f"id='{tableid}'")
    sync_condition=Helperdb.db_get_value('sys_table','sync_condition',f"id='{tableid}'")
    sync_order=Helperdb.db_get_value('sys_table','sync_order',f"id='{tableid}'")
    bixdata_fields=dict()
    rows=Helperdb.db_get('sys_field','*',f"tableid='{tableid}' AND sync_fieldid is not null AND sync_fieldid<>'' ")
    for row in rows:
        bixdata_fields[row['sync_fieldid']]=row['fieldid']
    if sync_condition:
        condition=sync_condition
    else:
        condition='true'

    if sync_order:
        order=f"ORDER BY {sync_order}"
    else:
        order=''
    sql=f"""
        SELECT *
        FROM {sync_table}
        WHERE {condition}
        {order}
    """
    syncrows=Helperdb.sql_query(sql)
    for syncrow in syncrows:
        sync_fields=dict()
        for key, field in syncrow.items():
            if key in bixdata_fields:
                sync_fields[bixdata_fields[key]]=field
    bixdata_sync_field=bixdata_fields[sync_field]
    print(sync_fields)


def download_attachment(request):
    recordid = request.POST.get('recordid')
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT filename from user_attachment WHERE recordid_='{recordid}'"
        )

        filename = cursor.fetchone()[0]

    file_path = os.path.join(base_dir, filename)


    with open(file_path, 'rb') as f:
        file_data = f.read()

def download_file(request, filename):
    # Costruisci il percorso completo al file
    filepath = os.path.join(settings.ATTACHMENTS_ROOT, filename)
    fs = FileSystemStorage(location='attachments') 
    filepath2=fs.location + "\\" + filename
    if os.path.exists(filepath2):
        # Restituisci il file come download
        return FileResponse(open(filepath2, 'rb'), as_attachment=True, filename=filename)
    else:
        raise Http404("File non trovato")


    response = HttpResponse(file_data, content_type='application/octet-stream')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def update_profile_pic(request):

    src = request.POST.get('image')

    src = src.split('static', 1)[1]

    base_path = 'bixdata_view/bixdata_app/static'
    image_path = os.path.join(base_path, src.strip(os.sep))

    destination_dir = 'bixdata_view/bixdata_app/static/images/users'

    image_path = 'bixdata_view/bixdata_app/static' + image_path





    userid=get_userid(request.user.id)

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM v_users WHERE sys_user_id='{userid}'"
        )

        user = dictfetchall(cursor)[0]

        username= user['username']

        fullname = user['first_name'] + ' ' + user['last_name']

        name = user['first_name'].lower()

        image1 = username + '.png'
        image2 = fullname + '.png'
        image3 = str(userid) + '.png'
        image4 = name + '.png'

        destination_path1 = os.path.join(destination_dir, image1)
        destination_path2 = os.path.join(destination_dir, image2)
        destination_path3 = os.path.join(destination_dir, image3)
        destination_path4 = os.path.join(destination_dir, image4)

        destination_path1 = os.path.normpath(destination_path1)
        destination_path2 = os.path.normpath(destination_path2)
        destination_path3 = os.path.normpath(destination_path3)
        destination_path4 = os.path.normpath(destination_path4)

        if os.path.exists(destination_path1):
            os.remove(destination_path1)

        if os.path.exists(destination_path2):
            os.remove(destination_path2)

        if os.path.exists(destination_path3):
            os.remove(destination_path3)

        if os.path.exists(destination_path4):
            os.remove(destination_path4)

        shutil.copy(image_path, destination_path1)
        shutil.copy(image_path, destination_path2)
        shutil.copy(image_path, destination_path3)
        shutil.copy(image_path, destination_path4)

    return JsonResponse({'success': True})

"""
User stats START
"""

def get_user_stats_page(request):
    context = dict()

    with connection.cursor() as cursor:
        cursor.execute(
            "SELECT * FROM v_users"
        )

        context['users'] = dictfetchall(cursor)

    return render(request, 'block/user/stats/statistics.html', context)


def get_user_stats_card(request):

    context = dict()
    context['date'] = datetime.datetime.now().strftime('%Y-%m-%d-%H:%M:%S')

    selected_users = request.POST.get('selectedUsers')
    selected_users = json.loads(selected_users)




    ids = [str(user) for user in selected_users]

    placeholders = ', '.join(['%s'] * len(ids))

    with connection.cursor() as cursor:
        cursor.execute(
            f"SELECT * FROM v_users WHERE sys_user_id IN ({placeholders})",
            ids
        )
        context['users'] = dictfetchall(cursor)

        block_list = dict()

        for user in context['users']:

            #dict contenente tutti i blocchi che verranno passati ad una pagina html che li cicla e li organizza per poi andare a caricarsi nella card
            #block_list['timesheet_chartblock'] = build_card_content(request, user['sys_user_id'])

            """
                esempio 
            
                intermediate_html = render_to_string('intermediate_template.html', {'block_list': block_list})

                # Ora puoi usare `intermediate_html` come una stringa nel template finale   
                context['intermediate_html'] = intermediate_html
                
                che poi viene passato al template user_stats_card
            """



            user['chartblock'] = build_card_content(request, user['sys_user_id'])
            sql = f"SELECT COUNT(recordid_) as opentasks FROM user_task WHERE status != 'Chiuso' AND user = {user['sys_user_id']}"
            user['opentasks'] = get_card_data(request,sql,user['sys_user_id'])[0]['opentasks']


            userid = user['sys_user_id']

            user_group = Helperdb.sql_query(f"SELECT groupid FROM sys_group_user WHERE userid = '{userid}'")

            if user_group:
                user_group = user_group[0]['groupid']

            if user_group == 60:
                user['taskblock'] = get_user_tasks_stats(request, userid)
            else:
                user['taskblock'] = ''



    return render(request, 'block/user/stats/stats_card.html', context)


def build_card_content(request, userid):

    block = []


    sql = f'SELECT SUM(totaltime_decimal),invoicestatus FROM user_timesheet WHERE date <= current_date AND date >= DATE_SUB(current_date,INTERVAL 14 day) and user={userid} GROUP BY invoicestatus'

    block = get_card_chart(request,userid,sql, 'donutchart', 'Invoiced Time')

    return block


def get_card_chart(request, userid, sql, type, chartname):

    with connection.cursor() as cursor2:
        cursor2.execute(sql)
        rows = cursor2.fetchall()
        formatted_rows = []
        for row in rows:
            formatted_row = [str(value) if not isinstance(value, (int, float)) else value for value in row]
            formatted_rows.append(formatted_row)

        rows = formatted_rows

        value = []
        labels = []

        for row in rows:
            value.append(row[0])
            labels.append(row[1])



        context = {
            'value': value,
            'labels': labels,
            'name': chartname,
            'id': uuid.uuid4().hex,
        }

        if type == 'barchart':
            return render_to_string('block/chart/barchart.html', context, request=request)
        elif type == 'piechart':
            return render_to_string('block/chart/piechart.html', context, request=request)
        elif type == 'linechart':
            return render_to_string('block/chart/linechart.html', context, request=request)
        elif type == 'horizontalbarchart':
            return render_to_string('block/chart/horizontalbarchart.html', context, request=request)
        elif type == 'donutchart':
            return render_to_string('block/chart_stats/donutchart.html', context, request=request)

def get_card_data(request, sql, userid):
    with connection.cursor() as cursor:
        cursor.execute(
            sql
        )

        rows = dictfetchall(cursor)

    return rows




"""
User Stats END
"""




def stampa_milestone(request):
    recordid = request.POST.get('recordid')

    tasks = Helperdb.sql_query(f"SELECT * FROM user_task WHERE recordidprojectmilestone_='{recordid}'")

    for task in tasks:
        fullname = Helperdb.sql_query(f"SELECT first_name, last_name FROM v_users WHERE sys_user_id = '{task['user']}'")
        fullname = fullname[0]['first_name'] + ' ' + fullname[0]['last_name']
        task['username'] = fullname


    milestone = Helperdb.sql_query(f"SELECT * FROM user_projectmilestone WHERE recordid_='{recordid}'")[0]


    context = dict()
    context['tasks'] = tasks
    context['milestone'] = milestone

    """
        filename = 'test.pdf'
        
        
        path = os.path.dirname(os.path.abspath(__file__))
        path = path.rsplit('views', 1)[0]
        filename_with_path = path + '\\static\\pdf\\' + filename
    
    
    
        script_dir = os.path.dirname(os.path.abspath(__file__))
        wkhtmltopdf_path = script_dir + '\\wkhtmltopdf.exe'
    
        config = pdfkit.configuration(wkhtmltopdf=wkhtmltopdf_path)
        content = render_to_string('pdf/milestone.html', context)
    
        pdfkit.from_string(content, filename_with_path, configuration=config)
    
        try:
            with open(filename_with_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/pdf")
                response['Content-Disposition'] = f'inline; filename={filename}'
    
            return response
    
        finally:
            os.remove(filename_with_path)
    """
    return render(request, 'pdf/milestone.html', context)


def get_user_tasks_stats(request, userid=''):

    request_type = ''
    if not userid:
        userid = get_userid(request.user.id)
        request_type = 'ajax'

    tasks = Helperdb.sql_query(f"SELECT * FROM user_task WHERE user='{userid}'")

    open_tasks = []
    closed_tasks = []
    expired_tasks = []


    for task in tasks:
        if task['status'] != 'Chiuso':
            open_tasks.append(task)
        if task['status'] == 'Chiuso':
            closed_tasks.append(task)
        if task['status'] == 'Scaduto':
            expired_tasks.append(task)


    open_tasks_count = len(open_tasks)
    closed_tasks_count = len(closed_tasks)
    expired_tasks_count = len(expired_tasks)





    closed_in_time = []
    days_diff= []


    for task in closed_tasks:
        if task['closedate'] and task['duedate']:
            if task['closedate'] < task['duedate']:
                closed_in_time.append(task)

            duedate = task['duedate']
            closedate = task['closedate']
            diff = (closedate - duedate).days
            days_diff.append(diff)

    closed_in_time = len(closed_in_time)

    if days_diff:
        avg_diff = round(sum(days_diff) / len(days_diff))
    else:
        avg_diff = 0

    data = dict()

    data = {
        'value': [open_tasks_count, closed_tasks_count, expired_tasks_count],
        'labels': ['Task aperti', 'Task chiusi', 'Task scaduti'],
        'name': 'tasks chart',
        'id': uuid.uuid4().hex,
    }

    tasks_chart = render_to_string('block/chart_stats/donutchart.html', data, request=request)


    context = dict()
    context['tasks'] = tasks
    context['open_tasks'] = open_tasks
    context['closed_tasks'] = closed_tasks
    context['expired_tasks'] = expired_tasks
    context['open_tasks_count'] = open_tasks_count
    context['closed_tasks_count'] = closed_tasks_count
    context['expired_tasks_count'] = expired_tasks_count
    context['closed_in_time'] = closed_in_time
    context['avg_close_day'] = avg_diff
    context['tasks_chart'] = tasks_chart

    if request_type == 'ajax':
        return JsonResponse({'tasks_stats': render_to_string('block/user/stats/tasks_stats.html', context)})
    else:
        return render_to_string('block/user/stats/tasks_stats.html', context, request=request)




def task_reminder(request):

    fields_dict = dict()

    recordid = request.POST.get('recordid')
    reminder_message = request.POST.get('reminderMessage')

    reminder_creator = get_userid(request.user.id)


    user = Helperdb.sql_query(f"SELECT * FROM v_users WHERE sys_user_id = {reminder_creator}")[0]
    fields_dict['reminder_creator'] = user['first_name'] + ' ' + user['last_name']


    fields_dict['task'] = Helperdb.sql_query(f"SELECT * FROM user_task WHERE recordid_ = {recordid}")[0]
    task_user = fields_dict['task']['user']

    fields_dict['email'] = Helperdb.sql_query(f"SELECT email from v_users WHERE sys_user_id = '{task_user}'")[0]['email']

    fields_dict['recordid'] = recordid
    fields_dict['reminder_message'] = reminder_message
    print(reminder_message)
    print(reminder_message)
    print(reminder_message)


    message = render_to_string('other/task_reminder.html', fields_dict)

    # return render(request, 'other/new_task.html', fields_dict)

    send_email(emails=[fields_dict['email']],
               subject=fields_dict['reminder_creator'] + ' Ti sollecita a svolgere questo task',
               html_message=message)

    return JsonResponse({'success': True})



def save_user_timesheet(request):

    timesheets = Helperdb.sql_query(f"select * from user_timesheet WHERE  true AND user_timesheet.recordid_ in(                       SELECT recordid_ FROM user_timesheet WHERE true AND ( invoicestatus like '%To invoice when%')                              ) AND date>'2024-01-01'")

    for timesheet in timesheets:
        request.POST = request.POST.copy()
        request.POST['recordid'] = timesheet['recordid_']
        request.POST['tableid'] = 'timesheet'
        request.POST['contextfunction'] = 'edit'
        save_record_fields(request)



    return JsonResponse({'success': True})


def logout_view(request):

    logout(request)
    return redirect('login')










